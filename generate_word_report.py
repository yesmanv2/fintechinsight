#!/usr/bin/env python3
"""生成微信支付客诉分析报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import json
from datetime import datetime, date
from collections import Counter, defaultdict

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_row(table, cells_data, bold=False, header=False, bg_color=None):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "2B579A")
        elif bg_color:
            set_cell_shading(cell, bg_color)
    return row

def create_styled_table(doc, headers, rows_data, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2B579A")
    
    # 添加数据行
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "E8EDF3")
    
    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    
    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    return table

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题 - 使用普通段落手动设置，避免内置heading样式在某些Word版本中显示不全"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(32)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(26)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x3D, 0x7E, 0xBB)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(22)
    return p

def add_paragraph_styled(doc, text, bold=False, italic=False, font_size=10.5, color=None, alignment=None, space_after=6):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_bullet_point(doc, text, bold_prefix="", indent_level=0):
    """添加要点列表"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.name = '微软雅黑'
        run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_callout_box(doc, text, bg_color="E8F0FE", border_color="2B579A"):
    """添加引用/高亮框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 使用左边框模拟callout效果
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    # 设置段落背景
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_horizontal_line(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def load_and_analyze_data():
    """加载数据并进行动态分析"""
    data_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "wxpay_complaints_raw.json")
    with open(data_path, 'r', encoding='utf-8') as f:
        all_posts = json.load(f)
    
    today = date.today()
    report_year = today.year
    report_month = today.month
    report_date_str = today.strftime('%Y年%m月%d日').replace('年0', '年').replace('月0', '月')
    
    # 计算总互动量
    total_interactions = 0
    for p in all_posts:
        for field in ['liked_count', 'comment_count', 'collected_count']:
            val = p.get(field, '0')
            try:
                total_interactions += int(str(val).replace(',', ''))
            except (ValueError, TypeError):
                pass
    
    # 月度统计
    monthly_counts = Counter()
    for p in all_posts:
        t = p.get('time', '')[:7]
        if t:
            monthly_counts[t] += 1
    
    # 时间范围
    all_times = sorted([p.get('time', '')[:10] for p in all_posts if p.get('time', '')])
    earliest = all_times[0] if all_times else '未知'
    latest = all_times[-1] if all_times else '未知'
    
    # 情绪分析
    sentiment_counts = Counter()
    for p in all_posts:
        s = p.get('sentiment', 'unknown')
        sentiment_counts[s] += 1
    
    # 分类统计
    category_counts = Counter()
    category_posts = defaultdict(list)
    for p in all_posts:
        cat = p.get('category', '其他')
        category_counts[cat] += 1
        category_posts[cat].append(p)
    
    # 关键词标签统计
    tag_counts = Counter()
    for p in all_posts:
        for tag in p.get('keyword_tags', []):
            tag_counts[tag] += 1
    
    # 当年帖子
    current_year_prefix = str(report_year)
    posts_current_year = sorted(
        [d for d in all_posts if d.get('time', '').startswith(current_year_prefix)],
        key=lambda x: x.get('time', '')
    )
    
    # 当月帖子
    current_month_prefix = f"{report_year}-{report_month:02d}"
    posts_current_month = [d for d in all_posts if d.get('time', '').startswith(current_month_prefix)]
    
    # 最高赞帖子 top10
    def get_likes(p):
        try:
            return int(str(p.get('liked_count', '0')).replace(',', ''))
        except (ValueError, TypeError):
            return 0
    
    top_posts = sorted(all_posts, key=get_likes, reverse=True)[:10]
    
    return {
        'all_posts': all_posts,
        'total_count': len(all_posts),
        'total_interactions': total_interactions,
        'monthly_counts': monthly_counts,
        'earliest': earliest,
        'latest': latest,
        'sentiment_counts': sentiment_counts,
        'category_counts': category_counts,
        'category_posts': category_posts,
        'tag_counts': tag_counts,
        'posts_current_year': posts_current_year,
        'posts_current_month': posts_current_month,
        'top_posts': top_posts,
        'report_year': report_year,
        'report_month': report_month,
        'report_date_str': report_date_str,
        'today': today,
    }


def generate_report():
    # ===== 动态加载数据 =====
    stats = load_and_analyze_data()
    today = stats['today']
    report_year = stats['report_year']
    report_month = stats['report_month']
    report_date_str = stats['report_date_str']
    total_count = stats['total_count']
    total_interactions = stats['total_interactions']
    monthly_counts = stats['monthly_counts']
    sentiment_counts = stats['sentiment_counts']
    category_counts = stats['category_counts']
    category_posts = stats['category_posts']
    tag_counts = stats['tag_counts']
    posts_current_year = stats['posts_current_year']
    posts_current_month = stats['posts_current_month']
    top_posts = stats['top_posts']
    all_posts = stats['all_posts']
    earliest = stats['earliest']
    latest = stats['latest']

    # 格式化互动量
    if total_interactions >= 10000:
        interaction_str = f"{total_interactions:,}次（约{total_interactions // 10000}万+次）"
    else:
        interaction_str = f"{total_interactions:,}次"
    
    # 格式化数据范围
    def format_date_cn(date_str):
        """将 2025-01-01 格式转为 2025年1月1日"""
        try:
            d = datetime.strptime(date_str, '%Y-%m-%d')
            return f"{d.year}年{d.month}月{d.day}日"
        except Exception:
            return date_str
    
    data_range_str = f"{format_date_cn(earliest)} — {format_date_cn(latest)}"

    doc = Document()
    
    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ===== 修改默认样式 =====
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    from docx.shared import Twips
    # 使用1.15倍行距（而非固定行距），这样不同字号的段落都不会被裁切
    style.paragraph_format.line_spacing_rule = None  # 清除固定行距
    style.paragraph_format.line_spacing = 1.15
    
    # ===== 封面区域 =====
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 主标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = Pt(42)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run("微信支付客诉分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.line_spacing = Pt(22)
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("基于小红书舆情数据的客诉话题分析")
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    add_horizontal_line(doc)
    
    # 信息区域（动态生成）
    info_items = [
        ("数据来源", "小红书舆情爬取数据"),
        ("分析时间", report_date_str),
        ("数据范围", data_range_str),
        ("样本量", f"{total_count}条微信支付客诉相关帖子"),
        ("总互动量", interaction_str),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}：")
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_label.font.name = '微软雅黑'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)
        run_val.font.name = '微软雅黑'
        run_val._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_val.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        p.paragraph_format.space_after = Pt(2)
    
    # ===== 分页 =====
    doc.add_page_break()
    
    # ===== 目录页 =====
    add_heading_styled(doc, "目  录", level=1)
    add_horizontal_line(doc)
    
    toc_items = [
        "一、总结",
        "二、六大核心问题分类与详细分析",
        "    2.1 自动续费 / 免密支付（占比30%）",
        "    2.2 风控误伤 / 账户冻结 / 封号（占比16%）",
        "    2.3 客服体验差 / 投诉无门（占比9%）",
        "    2.4 提现手续费（占比8%）",
        "    2.5 转账诈骗 / 资金追回（占比7%）",
        "    2.6 其他问题：分付/借贷、理财通",
        "三、痛点地图",
        "四、趋势分析",
        "五、建议与解决方案",
        "六、结论",
        f"附录：{report_year}年客诉帖子清单",
    ]
    for item in toc_items:
        is_sub = item.startswith("    ")
        p = doc.add_paragraph()
        run = p.add_run(item.strip())
        run.font.size = Pt(11 if not is_sub else 10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if not is_sub:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        else:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_after = Pt(4)
        if is_sub:
            p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_page_break()
    
    # ===== 一、总结 =====
    add_heading_styled(doc, "一、总结", level=1)
    add_horizontal_line(doc)
    
    add_callout_box(doc, 
        "2025年1月至今，微信支付在小红书上的客诉帖子共计186条，总互动量超过17万次，"
        "反映出用户对微信支付的不满高度集中在6个核心领域。",
        bg_color="E8F0FE", border_color="2B579A"
    )
    
    # 六大核心问题概览表
    add_paragraph_styled(doc, "六大核心问题概览：", bold=True, font_size=11, space_after=8)
    
    overview_table = create_styled_table(doc,
        ["排名", "问题类型", "帖子数", "占比", "最高单帖赞", "关键词"],
        [
            ["1", "自动续费/免密支付", "~55", "30%", "10,405", "入口太深、不知情扣费"],
            ["2", "风控误伤/封号", "~29", "16%", "12,449", "误判、解封慢、无申诉"],
            ["3", "客服体验差", "~17", "9%", "4,947", "找不到人工、AI循环"],
            ["4", "提现手续费", "~15", "8%", "5,857", "收费不合理、大额贵"],
            ["5", "转账被骗追回", "~13", "7%", "14,860", "被拉黑、维权难"],
            ["6", "先用后付/分付", "~5", "5%", "6,126", "关不掉、利率不透明"],
        ],
        col_widths=[1.5, 4, 1.5, 1.5, 2.5, 5]
    )
    
    doc.add_paragraph()
    
    # 关键发现
    add_paragraph_styled(doc, "关键发现：", bold=True, font_size=11)
    
    findings = [
        ("自动续费/免密支付是第一大痛点", '——超过50条帖子涉及，用户反复吐槽"不知道怎么关"、"被偷偷扣钱"。单篇帖子最高10405赞，说明这是一个广泛共鸣的普遍性问题。'),
        ("风控误伤是情绪最强烈的问题", '——"正常转账5000被风控审核3天"引发万级互动，用户愤怒值极高。'),
        ("找不到人工客服", '——"微信客服真的很难找"是最强共鸣帖之一（累计7000+赞），用户核心痛点是：人工客服入口隐蔽、打电话永远是AI、投诉后无人跟进。'),
        ("提现手续费让用户感到被割韭菜", "——虽然2025年8月有所下降，但用户对0.1%的费率仍然不满。"),
        ("转账安全需求巨大", '——"转账被拉黑如何追回"成为最高赞帖子（14860赞），说明大量用户有被骗后维权需求。'),
        ("先用后付被视为陷阱", '——6000+赞的帖子直指"先用后付有多坑"，用户认为关闭流程不透明。'),
    ]
    for prefix, detail in findings:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    
    doc.add_paragraph()
    add_callout_box(doc,
        "整体趋势：客诉量呈显著上升趋势。2026年Q1（1-3月）合计86条，已超过2025年全年的一半，"
        "3月单月47条创历史新高。",
        bg_color="FFF3E0", border_color="E65100"
    )
    
    doc.add_page_break()
    
    # ===== 二、详细分析 =====
    add_heading_styled(doc, "二、六大核心问题分类与详细分析", level=1)
    add_horizontal_line(doc)
    
    # ----- 2.1 自动续费 -----
    add_heading_styled(doc, "2.1 自动续费 / 免密支付（约55条，占比30%）", level=2)
    add_callout_box(doc, "核心痛点：操作路径隐蔽、用户不知情被扣费、关闭流程复杂")
    
    add_paragraph_styled(doc, "这是出现频率最高、用户共鸣最强烈的客诉类型。", bold=True, font_size=10.5)
    
    add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
    create_styled_table(doc,
        ["时间", "标题", "点赞", "评论"],
        [
            ["2025-07-21", "微信自动续费怎么取消？", "10,405", "2,444"],
            ["2025-03-12", "怎么关闭微信自动续费？", "4,547", "260"],
            ["2026-01-21", "微信自动续费怎么取消？", "2,960", "9"],
            ["2025-01-02", "微信关闭自动续费在哪里？自动续费关闭流程", "1,977", "63"],
            ["2025-11-18", "一起游刃有余~微信关闭自动续费", "1,830", "62"],
            ["2025-06-21", "怎么关闭自动续费", "1,552", "99"],
            ["2025-03-25", "关闭微信自动扣费", "1,472", "39"],
            ["2025-03-15", "💥又被扣钱❗自动续费忘记关啦❗❗", "110", "23"],
            ["2026-01-19", "自动续费套路大调查！你的钱怎么被无感扣走", "395", "18"],
            ["2025-05-27", "先用后付有多坑？", "6,126", "140"],
            ["2025-03-26", "多少人被免密支付坑了", "174", "25"],
        ],
        col_widths=[2.5, 8, 2, 2]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户核心诉求提炼：", bold=True, font_size=10.5)
    
    complaints_1 = [
        ('"在哪关？"', " — 关闭自动续费的入口太深，用户找不到路径"),
        ('"为什么偷偷扣钱？"', " — 很多用户不知道自己开通了自动续费，发现时已扣数月费用"),
        ('"先用后付关不掉"', " — 有用户反馈先用后付功能开了就关不掉，或关闭路径不清晰"),
        ('"免密支付太危险"', " — 用户担心免密支付的安全风险，但又不知道在哪里管理"),
        ('"被扣的钱能退吗？"', " — 大量帖子分享如何追回被自动续费扣的钱"),
    ]
    for prefix, detail in complaints_1:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    
    add_paragraph_styled(doc, '情绪特征：焦虑 + 愤怒 + 无助感。大量帖子都在教别人"怎么关"，说明这个操作路径的可发现性极差。', italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=10)
    
    add_horizontal_line(doc)
    
    # ----- 2.2 风控误伤 -----
    add_heading_styled(doc, "2.2 风控误伤 / 账户冻结 / 封号（29条，占比16%）", level=2)
    add_callout_box(doc, "核心痛点：正常使用被误判为风险、解封周期长、没有申诉通道")
    
    add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
    create_styled_table(doc,
        ["时间", "标题", "点赞", "评论"],
        [
            ["2025-01-24", "微信支付被风控了太无语", "12,449", "178"],
            ["2026-01-27", "微信解封值不值得用？", "6,091", "845"],
            ["2025-01-01", "微信支付被风控了太无语", "3,163", "787"],
            ["2025-11-27", "微信支付被风控了太无语", "2,139", "259"],
            ["2025-08-02", "微信有病吧 可以报警吗", "2,382", "762"],
            ["2025-11-21", "⚠️微信最近封号真的猛！", "1,002", "256"],
            ["2026-02-10", "微信被起诉", "431", "61"],
            ["2025-10-29", "平台乱封号还要报警笑死", "137", "64"],
            ["2026-03-16", "收到微信零钱突然被冻结，怎么办？", "3", "1"],
        ],
        col_widths=[2.5, 8, 2, 2]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户描述的典型场景：", bold=True, font_size=10.5)
    
    scenarios = [
        '"正常转账给朋友5000块，微信直接风控了说有风险，解除风控要审核3天"',
        '"永封几次了，微信被封麻木了"',
        '"微信怎么成天封号啊，真的无语了"',
        '"微信零钱被司法冻结"',
        '"企业微信吞我血汗钱！封禁账号不让提现"',
    ]
    for s in scenarios:
        add_bullet_point(doc, s)
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户核心诉求提炼：", bold=True, font_size=10.5)
    complaints_2 = [
        ('"正常操作为什么被风控？"', " — 风控模型误伤率高，用户感到莫名其妙"),
        ('"解封太慢"', " — 审核周期长，期间资金完全冻结影响生活"),
        ('"没有有效申诉渠道"', " — 用户不知道去哪申诉，或申诉了也没有反馈"),
        ('"永封后怎么提现？"', " — 封号后用户最关心资金安全，48小时提现窗口知晓率低"),
    ]
    for prefix, detail in complaints_2:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    
    add_horizontal_line(doc)
    
    # ----- 2.3 客服体验差 -----
    add_heading_styled(doc, "2.3 客服体验差 / 投诉无门（17条，占比9%）", level=2)
    add_callout_box(doc, "核心痛点：找不到人工客服、AI客服解决不了问题、投诉后无人跟进")
    
    add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
    create_styled_table(doc,
        ["时间", "标题", "点赞", "评论"],
        [
            ["2025-03-23", "聊聊微信投诉那些事", "4,947", "469"],
            ["2026-01-02", "微信客服真的很难找", "2,579", "327"],
            ["2025-12-20", "微信客服真的很难找", "1,249", "204"],
            ["2025-09-22", "微信投诉受理成功结果", "188", "111"],
            ["2025-06-26", "在12315投诉了腾讯的微信支付", "68", "11"],
            ["2025-05-11", "微信支付投诉根本没用", "11", "3"],
            ["2026-03-07", "怎么投诉微信支付啊", "2", "0"],
        ],
        col_widths=[2.5, 8, 2, 2]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户核心吐槽：", bold=True, font_size=10.5)
    complaints_3 = [
        '"想联系微信人工客服处理支付问题，入口藏得太深了，打电话也一直是AI"',
        '"微信支付投诉根本没用"',
        '"在12315投诉了腾讯的微信支付" — 用户不得不求助外部监管',
        '"投诉受理成功结果" — 用户发帖庆祝投诉成功，说明常规投诉渠道极难成功',
    ]
    for c in complaints_3:
        add_bullet_point(doc, c)
    
    add_paragraph_styled(doc, "情绪特征：极度挫败。用户在遇到风控、扣费等问题后，发现无法通过客服解决，愤怒叠加。", italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=10)
    
    add_horizontal_line(doc)
    
    # ----- 2.4 提现手续费 -----
    add_heading_styled(doc, "2.4 提现手续费（15条，占比8%）", level=2)
    add_callout_box(doc, '核心痛点：0.1%手续费被视为"不合理"、大额提现成本高')
    
    add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
    create_styled_table(doc,
        ["时间", "标题", "点赞", "评论"],
        [
            ["2025-06-15", "微信提现避坑指南｜这六种真的免费！", "5,857", "346"],
            ["2025-07-31", "我的天！微信提现手续费下限降了！", "667", "72"],
            ["2026-02-19", "微信提现2万要20块手续费？！！！！！", "493", "337"],
            ["2025-06-17", "微信提现免手续费，适用于不急用钱的人", "110", "11"],
            ["2025-07-28", "WX零钱通提现霸王条款！", "43", "33"],
        ],
        col_widths=[2.5, 8, 2, 2]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户核心诉求：", bold=True, font_size=10.5)
    complaints_4 = [
        ('"为什么提现要收费？钱是我的啊！"', " — 用户认为把自己的钱提到自己的银行卡不应收费"),
        ('"大额提现手续费太高"', " — 2万提现20块，用户接受不了"),
        ('"零钱通提现也收费？"', ' — 用户觉得这是"霸王条款"'),
    ]
    for prefix, detail in complaints_4:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    add_paragraph_styled(doc, "正面：2025年8月手续费下降受到好评，说明用户对改善是有正面回应的。", italic=True, color=RGBColor(0x00, 0x7A, 0x33), font_size=10)
    
    add_horizontal_line(doc)
    
    # ----- 2.5 转账诈骗 -----
    add_heading_styled(doc, "2.5 转账诈骗 / 资金追回（13条，占比7%）", level=2)
    add_callout_box(doc, "核心痛点：转账后被拉黑无法追回、维权路径不清晰")
    
    add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
    create_styled_table(doc,
        ["时间", "标题", "点赞", "评论"],
        [
            ["2025-04-14", "微信转账被拉黑后我的操作，已全额追回！", "14,860", "1,334"],
            ["2026-02-26", "转账之后被拉黑，其实简单一招就能拦截挽回", "1,040", "14"],
            ["2025-12-18", "别让付款码，成了骗子的提款码！", "241", "11"],
            ["2026-03-05", "被诈骗微信付款怎么追回？3步教你合法维权", "4", "0"],
            ["2025-03-28", "天塌了，睡醒发现亲属卡成为诈骗工具了", "26", "4"],
            ["2026-01-03", "新型微信支付诈骗", "17", "31"],
        ],
        col_widths=[2.5, 8, 2, 2]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "用户核心诉求：", bold=True, font_size=10.5)
    complaints_5 = [
        ('"被骗了钱怎么追回？"', " — 最高赞帖子（14860赞）就是分享追回经验，说明需求巨大"),
        ('"亲属卡被利用为诈骗工具"', " — 产品设计的安全漏洞"),
        ('"希望微信能拦截可疑转账"', " — 用户期望更强的实时风险拦截"),
    ]
    for prefix, detail in complaints_5:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    
    add_horizontal_line(doc)
    
    # ----- 2.6 其他问题 -----
    add_heading_styled(doc, "2.6 其他问题：分付/借贷、理财通", level=2)
    
    add_paragraph_styled(doc, "分付/借贷问题：", bold=True, font_size=10.5)
    add_bullet_point(doc, '"我嘞个豆，微信分付巨坑！"（37赞）')
    add_bullet_point(doc, '"微信分付大骗子"')
    add_paragraph_styled(doc, '核心问题：用户对分付的利率、还款规则、以及"是否值得开通"有疑虑。', font_size=10)
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "理财通相关投诉：", bold=True, font_size=10.5)
    add_bullet_point(doc, '"微信理财通 你是真能骗"（690赞）')
    add_bullet_point(doc, '"腾讯理财通诈骗犯"')
    add_bullet_point(doc, '"关闭微信\'腾讯微保\'骗局"')
    add_paragraph_styled(doc, "核心问题：用户对理财通产品信任度低，同时大量假冒微保的诈骗困扰用户。", font_size=10)
    
    doc.add_page_break()
    
    # ===== 三、痛点地图 =====
    add_heading_styled(doc, "三、痛点地图", level=1)
    add_horizontal_line(doc)
    
    add_paragraph_styled(doc, "按影响面 × 情绪强度排列：", bold=True, font_size=11)
    
    pain_map_table = create_styled_table(doc,
        ["痛点等级", "问题", "帖子数", "特征"],
        [
            ["🔴 极高", "自动续费/免密支付", "~55", "影响面最广、赞数最高、长期未解决"],
            ["🔴 极高", "风控误伤/封号", "~29", "情绪最强烈、愤怒值最高"],
            ["🟠 高", "客服找不到", "~17", "高共鸣、与其他问题叠加放大不满"],
            ["🟠 高", "提现手续费", "~15", "中等愤怒、但涉及真金白银"],
            ["🟡 中", "转账被骗追回", "~13", "单帖赞数最高（14860）、需求刚性"],
            ["🟡 中", "先用后付/分付", "~5", "高愤怒但影响面较小"],
            ["🟢 较低", "理财通/微保", "~4", "信任问题 + 被诈骗利用"],
        ],
        col_widths=[2.5, 4, 2, 7.5]
    )
    
    doc.add_paragraph()
    
    # ===== 四、趋势分析 =====
    add_heading_styled(doc, "四、趋势分析", level=1)
    add_horizontal_line(doc)
    
    add_paragraph_styled(doc, "月度客诉量趋势：", bold=True, font_size=11)
    
    trend_data = [
        ["2025年1月", "7", "█" * 7],
        ["2025年2月", "1", "█"],
        ["2025年3月", "14", "█" * 14],
        ["2025年4月", "5", "█" * 5],
        ["2025年5月", "6", "█" * 6],
        ["2025年6月", "7", "█" * 7],
        ["2025年7月", "14", "█" * 14],
        ["2025年8月", "8", "█" * 8],
        ["2025年9月", "7", "█" * 7],
        ["2025年10月", "8", "█" * 8],
        ["2025年11月", "10", "█" * 10],
        ["2025年12月", "13", "█" * 13],
        ["2026年1月", "19", "█" * 19],
        ["2026年2月", "20", "█" * 20],
        ["2026年3月", "47 ⬆", "█" * 47],
    ]
    
    trend_table = create_styled_table(doc,
        ["月份", "数量", "趋势"],
        trend_data,
        col_widths=[3, 2, 11]
    )
    
    doc.add_paragraph()
    add_paragraph_styled(doc, "关键发现：", bold=True, font_size=11)
    
    trend_findings = [
        ("2026年Q1客诉量暴增", " — 仅3月上半月就有47条，可能与年初支付新规执行、消费者维权意识增强有关"),
        ("2025年Q4是封号投诉高峰期", ' — 与"清朗行动"相关'),
        ("自动续费类帖子全年均匀分布", " — 说明是一个长期未解决的结构性问题"),
    ]
    for prefix, detail in trend_findings:
        add_bullet_point(doc, detail, bold_prefix=prefix)
    
    doc.add_page_break()
    
    # ===== 五、建议与解决方案 =====
    add_heading_styled(doc, "五、建议与解决方案", level=1)
    add_horizontal_line(doc)
    
    # 5.1
    add_heading_styled(doc, "5.1 自动续费 / 免密支付管理（优先级：🔴 最高）", level=2)
    create_styled_table(doc,
        ["建议", "具体方案"],
        [
            ["优化入口可发现性", "在「微信支付」主页增加醒目的「扣费管理」入口，而非埋在三级菜单里"],
            ["主动提醒机制", "每月推送一次「本月自动续费汇总」通知，让用户知道被扣了什么钱"],
            ["一键管理页", "建立一个统一的「我的自动续费/免密支付」管理页面，所有代扣服务一目了然"],
            ["扣费前确认", "对长期未使用但仍在自动续费的服务，在扣费前发送确认通知"],
            ["先用后付优化", "先用后付的开通/关闭流程需要更透明，增加关闭引导"],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_paragraph()
    
    # 5.2
    add_heading_styled(doc, "5.2 风控策略优化（优先级：🔴 最高）", level=2)
    create_styled_table(doc,
        ["建议", "具体方案"],
        [
            ["降低误伤率", "优化风控模型，对正常用户（高信用分、长期用户）的小额转账放宽阈值"],
            ["缩短审核周期", '将风控审核从"3天"缩短到"2小时快速通道"，至少对低风险case'],
            ["透明化", '告诉用户具体为什么被风控，而不只是笼统的"有风险"'],
            ["自助解封", "提供人脸+短信验证的快速自助解封通道，减少等待"],
            ["封号后资金保障", "更主动提醒用户48小时提现窗口，而不是让用户自己上网搜"],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_paragraph()
    
    # 5.3
    add_heading_styled(doc, "5.3 客服体验升级（优先级：🟠 高）", level=2)
    create_styled_table(doc,
        ["建议", "具体方案"],
        [
            ["人工客服入口外露", "在支付问题场景（如风控、扣费异议）中，直接展示人工客服按钮"],
            ["AI客服增加转人工选项", "当AI无法解决时，快速转接人工，而不是循环播放"],
            ["投诉工单跟踪", "建立可追踪的工单系统，用户可以查看处理进度"],
            ["95017改善", "优化电话客服体验，减少AI拦截层数"],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_paragraph()
    
    # 5.4
    add_heading_styled(doc, "5.4 提现手续费（优先级：🟡 中）", level=2)
    create_styled_table(doc,
        ["建议", "具体方案"],
        [
            ["继续降低费率", "2025年8月已下降，建议继续降低或对小额提现免费"],
            ["免费额度提升", "增加每月免费提现额度（目前1000元太低）"],
            ["清晰告知", "在提现页面明确展示手续费金额和免费替代方案"],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_paragraph()
    
    # 5.5
    add_heading_styled(doc, "5.5 转账安全 / 防诈骗（优先级：🟡 中）", level=2)
    create_styled_table(doc,
        ["建议", "具体方案"],
        [
            ["转账后延时到账", "对大额转账到非好友/新好友，默认24小时到账，给用户反悔窗口"],
            ["被拉黑后冻结", "转账后被拉黑，自动冻结该笔资金并通知用户"],
            ["亲属卡安全", "亲属卡增加单笔/单日限额设置，异常消费实时通知"],
            ["维权指南", '在App内提供官方的"被骗后怎么办"指引，减少用户去小红书求助'],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_page_break()
    
    # ===== 六、结论 =====
    add_heading_styled(doc, "六、结论", level=1)
    add_horizontal_line(doc)
    
    add_paragraph_styled(doc, "微信支付面临的客诉问题可以归结为两个本质矛盾：", bold=True, font_size=11)
    
    doc.add_paragraph()
    
    # 矛盾1
    add_heading_styled(doc, "矛盾一：安全与体验的矛盾", level=3)
    add_paragraph_styled(doc, 
        '风控越严，误伤越多；风控放松，诈骗增加。用户期望的是"精准的安全保障"，而不是"一刀切的限制"。'
        '29条风控相关帖子和13条转账诈骗帖子，从两个方向同时挤压着微信支付的决策空间。')
    
    # 矛盾2
    add_heading_styled(doc, "矛盾二：商业化与用户信任的矛盾", level=3)
    add_paragraph_styled(doc,
        '自动续费/免密支付为商户带来GMV，但不透明的操作路径损害了用户信任。提现手续费是收入来源，'
        '但用户觉得"我的钱凭什么要交费"。55条自动续费帖子和15条提现手续费帖子，说明商业模式正在透支用户耐心。')
    
    doc.add_paragraph()
    
    add_callout_box(doc,
        '最核心的改进方向：提升透明度。无论是风控原因、扣费明细、还是客服进度，'
        '用户要的不是"完美"，而是"知情权"。让用户看得见、关得掉、找得到人 —— '
        '这三个"得"，就是解决80%客诉的钥匙。',
        bg_color="E8F5E9", border_color="1B5E20"
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ===== 附录：当年客诉帖子清单 =====
    doc.add_page_break()
    
    # 动态生成附录标题
    year_posts = posts_current_year
    if year_posts:
        year_months = sorted(set(p.get('time', '')[:7] for p in year_posts))
        first_month = year_months[0] if year_months else f"{report_year}-01"
        last_month = year_months[-1] if year_months else f"{report_year}-{report_month:02d}"
        
        def month_key_to_cn(mk):
            parts = mk.split('-')
            return f"{parts[0]}年{int(parts[1])}月"
        
        appendix_title = f"附录：{month_key_to_cn(first_month)}-{int(last_month.split('-')[1])}月客诉帖子清单"
    else:
        appendix_title = f"附录：{report_year}年客诉帖子清单"
    
    add_heading_styled(doc, appendix_title, level=1)
    add_horizontal_line(doc)
    
    add_paragraph_styled(doc, 
        f"以下为{report_year}年期间，小红书上微信支付相关客诉帖子的完整清单（共{len(year_posts)}条），"
        "按时间顺序排列。每条帖子均附有原始链接，可点击查看原文。",
        font_size=10, space_after=10)
    
    # 按月份分组输出
    months_map = {}
    for post in year_posts:
        month_key = post.get('time', '')[:7]
        if month_key not in months_map:
            months_map[month_key] = []
        months_map[month_key].append(post)
    
    for month_key in sorted(months_map.keys()):
        month_posts = months_map[month_key]
        parts = month_key.split('-')
        label = f"{parts[0]}年{int(parts[1])}月"
        
        add_heading_styled(doc, f"{label}（{len(month_posts)}条）", level=2)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        headers_text = ["序号", "标题", "日期"]
        for i, h in enumerate(headers_text):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "2B579A")
        
        # 数据行
        for idx, post in enumerate(month_posts):
            row = table.add_row()
            
            # 序号
            cell0 = row.cells[0]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(str(idx + 1))
            run0.font.size = Pt(8)
            run0.font.name = '微软雅黑'
            run0._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 标题（带超链接）
            cell1 = row.cells[1]
            p1 = cell1.paragraphs[0]
            title_text = post.get('title', '')
            link = post.get('link', '')
            if link:
                # 添加超链接
                from docx.oxml import OxmlElement
                hyperlink = OxmlElement('w:hyperlink')
                r_id = doc.part.relate_to(link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), '微软雅黑')
                rFonts.set(qn('w:eastAsia'), '微软雅黑')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '16')  # 8pt = 16 half-points
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '16')
                rPr.append(szCs)
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), '2B579A')
                rPr.append(color_el)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                new_run.append(rPr)
                new_run.text = title_text
                hyperlink.append(new_run)
                p1._element.append(hyperlink)
            else:
                run1 = p1.add_run(title_text)
                run1.font.size = Pt(8)
                run1.font.name = '微软雅黑'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 日期
            cell2 = row.cells[2]
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(post.get('time', '')[:10])
            run2.font.size = Pt(8)
            run2.font.name = '微软雅黑'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 隔行背景
            if idx % 2 == 1:
                for c in row.cells:
                    set_cell_shading(c, "E8EDF3")
        
        # 设置列宽
        for row in table.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(11)
            row.cells[2].width = Cm(3.5)
        
        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)
        
        doc.add_paragraph()
    
    add_paragraph_styled(doc, 
        f"共计 {len(year_posts)} 条帖子。所有链接均指向小红书原帖，可直接点击查看。",
        italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=9, space_after=12)

    # 页脚信息
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_line.add_run("— 报告结束 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_time_str = today.strftime('%Y-%m-%d')
    run = footer_info.add_run(f"报告生成时间：{gen_time_str} | 数据源：小红书舆情爬取系统")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ===== 保存（动态文件名） =====
    output_filename = f"微信支付客诉分析报告_{report_year}年{report_month:02d}月.docx"
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), output_filename)
    doc.save(output_path)
    print(f"✅ Word报告已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    generate_report()
