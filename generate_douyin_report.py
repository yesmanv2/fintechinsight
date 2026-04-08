#!/usr/bin/env python3
"""
抖音舆情分析报告生成器
从 douyin_data.json 读取数据，生成独立 Word 分析报告。

Usage:
    python3 generate_douyin_report.py
    python3 generate_douyin_report.py --year 2026 --month 4
    python3 generate_douyin_report.py --output-dir ./reports
"""

import os
import sys
import json
import argparse
from datetime import datetime, date
from collections import Counter, defaultdict

# 复用现有样式工具函数
from generate_word_report import (
    set_cell_shading,
    create_styled_table,
    add_heading_styled,
    add_paragraph_styled,
    add_bullet_point,
    add_callout_box,
    add_horizontal_line,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'douyin_data.json')

BRAND_COLOR = RGBColor(0xFE, 0x2C, 0x55)  # 抖音品牌色
BRAND_HEX = "FE2C55"


def format_number(n):
    if isinstance(n, str):
        try: n = int(n)
        except: return n
    if n >= 10000:
        return f"{n / 10000:.1f}万"
    return f"{n:,}"


def get_interactions(video):
    try:
        like = int(video.get('like_count', 0) or 0)
        comment = int(video.get('comment_count', 0) or 0)
        share = int(video.get('share_count', 0) or 0)
        return like + comment + share
    except (ValueError, TypeError):
        return 0


def load_data():
    if not os.path.exists(DATA_FILE):
        print(f"❌ 数据文件不存在: {DATA_FILE}")
        print("💡 请先运行: python3 crawl_douyin.py")
        sys.exit(1)
    
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        raw = json.load(f)
    
    return raw.get('videos', []), raw.get('meta', {})


def generate_report(year=None, month=None, output_dir=None):
    """生成抖音舆情分析报告"""
    videos, meta = load_data()
    
    if not videos:
        print("❌ 没有数据")
        return None
    
    today = date.today()
    report_year = year or today.year
    report_month = month or today.month
    report_date_str = today.strftime('%Y年%m月%d日').replace('年0', '年').replace('月0', '月')
    
    # 数据分析
    total = len(videos)
    search_videos = [v for v in videos if v.get('source') == 'search' or v.get('source') == 'search_dom']
    hot_videos = [v for v in videos if v.get('source') == 'hot_topic']
    
    # 情绪分析
    sentiment_stats = Counter(v.get('sentiment', 'neutral') for v in videos)
    
    # 分类统计
    category_stats = Counter()
    for v in videos:
        for cat in v.get('categories', ['其他']):
            category_stats[cat] += 1
    
    # 关键词统计
    keyword_stats = Counter(v.get('search_keyword', '') for v in search_videos)
    
    # Top 互动视频
    top_videos = sorted(videos, key=get_interactions, reverse=True)[:15]
    
    # 负面视频
    negative_videos = [v for v in videos if v.get('sentiment') == 'negative']
    neg_sorted = sorted(negative_videos, key=get_interactions, reverse=True)[:10]
    
    # 月度统计
    monthly_stats = Counter()
    for v in videos:
        t = v.get('publish_time', '')[:7]
        if t:
            monthly_stats[t] += 1
    
    total_interactions = sum(get_interactions(v) for v in videos)
    
    # ==================== 创建文档 ====================
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = Pt(42)
    run = title_p.add_run("抖音舆情分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于抖音网页端的支付/金融短视频舆情监控")
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    add_horizontal_line(doc)
    
    info_items = [
        ("数据来源", "抖音网页端直接爬取"),
        ("分析时间", report_date_str),
        ("总样本量", f"{total} 条短视频"),
        ("搜索视频", f"{len(search_videos)} 条"),
        ("热搜话题", f"{len(hot_videos)} 条"),
        ("总互动量", format_number(total_interactions)),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}：")
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        run_l.font.name = '微软雅黑'
        run_l._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.name = '微软雅黑'
        run_v._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_v.font.color.rgb = BRAND_COLOR
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ==================== 一、概览 ====================
    add_heading_styled(doc, "一、数据概览", level=1)
    add_horizontal_line(doc)
    
    neg_count = sentiment_stats.get('negative', 0)
    neg_pct = f"{neg_count / total * 100:.1f}%" if total > 0 else "0%"
    
    add_callout_box(doc,
        f"本次共采集 {total} 条抖音短视频数据，其中搜索采集 {len(search_videos)} 条，"
        f"热搜话题 {len(hot_videos)} 条。负面情绪视频占比 {neg_pct}，"
        f"总互动量超过 {format_number(total_interactions)} 次。",
        bg_color="FFF0F0", border_color=BRAND_HEX)
    
    # 分类统计表
    add_paragraph_styled(doc, "内容分类统计：", bold=True, font_size=11, space_after=8)
    cat_rows = []
    for cat, cnt in category_stats.most_common():
        pct = f"{cnt / total * 100:.1f}%" if total > 0 else "0%"
        cat_rows.append([cat, str(cnt), pct])
    create_styled_table(doc, ["分类", "视频数", "占比"], cat_rows, col_widths=[4, 3, 3])
    
    doc.add_paragraph()
    
    # 情绪分布
    add_paragraph_styled(doc, "情绪分布：", bold=True, font_size=11, space_after=8)
    sentiment_labels = {'positive': '正面', 'negative': '负面', 'neutral': '中性'}
    sent_rows = []
    for key in ['positive', 'neutral', 'negative']:
        cnt = sentiment_stats.get(key, 0)
        pct = f"{cnt / total * 100:.1f}%" if total > 0 else "0%"
        bar = "█" * max(1, int(cnt / total * 30)) if total > 0 else "█"
        sent_rows.append([sentiment_labels.get(key, key), str(cnt), pct, bar])
    create_styled_table(doc, ["情绪", "数量", "占比", "分布"], sent_rows, col_widths=[2, 2, 2, 10])
    
    doc.add_page_break()
    
    # ==================== 二、关键词分析 ====================
    add_heading_styled(doc, "二、关键词采集分析", level=1)
    add_horizontal_line(doc)
    
    add_paragraph_styled(doc, "各搜索关键词的采集结果：", bold=True, font_size=11, space_after=8)
    kw_rows = []
    for kw, cnt in keyword_stats.most_common(20):
        if kw:
            kw_videos = [v for v in search_videos if v.get('search_keyword') == kw]
            kw_interactions = sum(get_interactions(v) for v in kw_videos)
            kw_neg = sum(1 for v in kw_videos if v.get('sentiment') == 'negative')
            kw_rows.append([kw, str(cnt), format_number(kw_interactions),
                          f"{kw_neg / cnt * 100:.0f}%" if cnt > 0 else "0%"])
    if kw_rows:
        create_styled_table(doc, ["关键词", "视频数", "总互动", "负面率"], kw_rows, col_widths=[5, 2, 3, 2])
    
    doc.add_page_break()
    
    # ==================== 三、热门视频 ====================
    add_heading_styled(doc, "三、最高互动视频 Top 15", level=1)
    add_horizontal_line(doc)
    
    hot_rows = []
    for v in top_videos:
        title = v.get('title', '')[:40]
        author = v.get('author', '')[:10]
        likes = format_number(v.get('like_count', 0))
        comments = format_number(v.get('comment_count', 0))
        plays = format_number(v.get('play_count', 0))
        sent = sentiment_labels.get(v.get('sentiment', 'neutral'), '中性')
        hot_rows.append([title, author, likes, comments, plays, sent])
    
    if hot_rows:
        create_styled_table(doc,
            ["标题", "作者", "点赞", "评论", "播放", "情绪"],
            hot_rows,
            col_widths=[6, 2.5, 1.5, 1.5, 2, 1.5])
    
    doc.add_page_break()
    
    # ==================== 四、负面舆情 ====================
    add_heading_styled(doc, "四、负面舆情分析", level=1)
    add_horizontal_line(doc)
    
    add_callout_box(doc,
        f"共发现 {len(negative_videos)} 条负面情绪视频，占全部视频的 {neg_pct}。"
        f"以下为互动量最高的负面视频，需重点关注。",
        bg_color="FFF3E0", border_color="E65100")
    
    doc.add_paragraph()
    
    neg_rows = []
    for v in neg_sorted:
        title = v.get('title', '')[:45]
        likes = format_number(v.get('like_count', 0))
        time_str = v.get('publish_time', '')[:10]
        neg_rows.append([time_str, title, likes])
    
    if neg_rows:
        add_paragraph_styled(doc, "负面高互动视频 Top 10：", bold=True, font_size=11, space_after=8)
        create_styled_table(doc, ["时间", "标题", "点赞"], neg_rows, col_widths=[2.5, 10, 2])
    
    # 负面主题聚类
    doc.add_paragraph()
    add_paragraph_styled(doc, "负面内容主题分布：", bold=True, font_size=11)
    neg_cats = Counter()
    for v in negative_videos:
        for cat in v.get('categories', ['其他']):
            neg_cats[cat] += 1
    for cat, cnt in neg_cats.most_common():
        pct = f"{cnt / len(negative_videos) * 100:.0f}%" if negative_videos else "0%"
        add_bullet_point(doc, f"：{cnt} 条（{pct}）", bold_prefix=cat)
    
    doc.add_page_break()
    
    # ==================== 五、热搜话题 ====================
    if hot_videos:
        add_heading_styled(doc, "五、热搜关联话题", level=1)
        add_horizontal_line(doc)
        
        add_paragraph_styled(doc, "抖音热搜中与支付/金融相关的话题：", bold=True, font_size=11, space_after=8)
        
        hot_topic_rows = []
        for v in hot_videos:
            title = v.get('title', '')
            hot_value = format_number(v.get('like_count', 0))
            cats = ', '.join(v.get('categories', []))
            hot_topic_rows.append([title, hot_value, cats])
        
        if hot_topic_rows:
            create_styled_table(doc, ["话题", "热度", "分类"], hot_topic_rows, col_widths=[8, 3, 3])
        
        doc.add_page_break()
    
    # ==================== 六、月度趋势 ====================
    if monthly_stats:
        add_heading_styled(doc, "六、月度趋势", level=1)
        add_horizontal_line(doc)
        
        sorted_months = sorted(monthly_stats.keys())
        if sorted_months:
            max_count = max(monthly_stats.values())
            trend_rows = []
            for mk in sorted_months:
                cnt = monthly_stats[mk]
                parts = mk.split('-')
                label = f"{parts[0]}年{int(parts[1])}月" if len(parts) == 2 else mk
                bar_len = max(1, int(cnt / max_count * 30))
                trend_rows.append([label, str(cnt), "█" * bar_len])
            create_styled_table(doc, ["月份", "视频数", "趋势"], trend_rows, col_widths=[3, 2, 11])
    
    doc.add_paragraph()
    
    # ==================== 结论 ====================
    add_heading_styled(doc, "结论与建议", level=1)
    add_horizontal_line(doc)
    
    top_cats = category_stats.most_common(3)
    if top_cats:
        cat_names = "、".join(f"「{c[0]}」" for c in top_cats)
        add_paragraph_styled(doc,
            f"抖音平台上与支付/金融相关的短视频内容主要集中在 {cat_names} 三个领域。",
            bold=True, font_size=11)
    
    doc.add_paragraph()
    
    add_callout_box(doc,
        "建议定期监控抖音平台的舆情变化，特别关注负面情绪视频的传播趋势。"
        "对于高互动量的负面视频，建议及时了解用户痛点并做出响应。",
        bg_color="E8F5E9", border_color="1B5E20")
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 报告结束 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"报告生成时间：{today.strftime('%Y-%m-%d')} | 数据源：抖音网页端爬取")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 保存
    if output_dir is None:
        output_dir = BASE_DIR
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"抖音舆情分析报告_{report_year}年{report_month:02d}月.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"✅ 报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="抖音舆情分析报告生成器")
    parser.add_argument('--year', type=int, default=None)
    parser.add_argument('--month', type=int, default=None)
    parser.add_argument('--output-dir', type=str, default=None)
    args = parser.parse_args()
    
    generate_report(year=args.year, month=args.month, output_dir=args.output_dir)
