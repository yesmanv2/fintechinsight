#!/usr/bin/env python3
"""Read the markdown report and convert it to a styled Word document."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

_dir = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(_dir, '\u5fae\u4fe1\u652f\u4ed8\u7b2c\u4e00\u9009\u62e9\u6218\u7565\u7814\u7a76\u62a5\u544a.md')


def set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, header_row, data_rows, col_widths=None):
    ncols = len(header_row)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, 9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, '1A3C6E')
    # data
    for idx, row_data in enumerate(data_rows):
        row = table.add_row()
        for i in range(ncols):
            cell = row.cells[i]
            text = row_data[i] if i < len(row_data) else ''
            run = cell.paragraphs[0].add_run(text)
            set_font(run, 9)
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'E8EDF3')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    # borders
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    ))
    return table


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        set_font(run, 20, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_font(run, 15, bold=True, color=RGBColor(0x2B, 0x57, 0x9A))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_font(run, 13, bold=True, color=RGBColor(0x3D, 0x7E, 0xBB))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    # Handle **bold** markers inside text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, size, bold=True, color=color)
        else:
            run = p.add_run(part)
            set_font(run, size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def add_callout(doc, text, bg='E8F0FE', border='2B579A'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{border}"/>'
        f'</w:pBdr>'
    ))
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg}"/>'
    ))
    # Handle **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, 10, bold=True, color=RGBColor(0x33, 0x33, 0x33))
            run.italic = True
        else:
            run = p.add_run(part)
            set_font(run, 10, italic=True, color=RGBColor(0x33, 0x33, 0x33))
    p.paragraph_format.line_spacing = Pt(18)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    ))


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, 10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, 10)
    p.paragraph_format.space_after = Pt(3)


# =============== MARKDOWN PARSER ===============

def parse_md_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # skip separator row
        if all(set(c) <= set('-: ') for c in cells):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def process_md(md_text, doc):
    lines = md_text.split('\n')
    i = 0
    in_blockquote = False
    blockquote_buf = []

    def flush_bq():
        nonlocal blockquote_buf
        if blockquote_buf:
            text = '\n'.join(blockquote_buf)
            # Choose callout style based on content
            if any(k in text for k in ['\u5371\u9669', '\u8680\u98df', '\u53e0\u52a0\u6548\u5e94']):
                add_callout(doc, text, bg='FFF3E0', border='E65100')
            elif '\u6ce8\uff1a' in text:
                add_callout(doc, text, bg='FFF8E1', border='F9A825')
            elif '\u7ed3\u8bed' in text:
                add_callout(doc, text, bg='E8F5E9', border='1B5E20')
            else:
                add_callout(doc, text)
            blockquote_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            flush_bq()
            in_blockquote = False
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            content = stripped[2:]
            # remove leading **...**: pattern for bold prefix
            blockquote_buf.append(content)
            in_blockquote = True
            i += 1
            continue
        elif in_blockquote:
            flush_bq()
            in_blockquote = False

        # Heading
        m = re.match(r'^(#{1,3})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            add_heading(doc, text, level)
            add_hr(doc)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                # clean bold markers from table cells
                clean_h = [re.sub(r'\*\*(.*?)\*\*', r'\1', h) for h in headers]
                clean_r = [[re.sub(r'\*\*(.*?)\*\*', r'\1', c) for c in row] for row in rows]
                add_table(doc, clean_h, clean_r)
                doc.add_paragraph()
            continue

        # Bullet list
        if stripped.startswith('- '):
            text = stripped[2:]
            add_bullet(doc, text)
            i += 1
            continue

        # Numbered list
        nm = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if nm:
            text = nm.group(2)
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', f"{nm.group(1)}. {text}")
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, 10, bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, 10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        # TOC link lines (skip - [...])
        if stripped.startswith('- ['):
            i += 1
            continue

        # Regular paragraph
        add_para(doc, stripped)
        i += 1

    flush_bq()


def generate():
    # Read markdown
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        md = f.read()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    for margin in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section, margin, Cm(2.5))

    # Default style
    style = doc.styles['Normal']
    style.font.name = '\u5fae\u8f6f\u96c5\u9ed1'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')

    # Split into sections by top-level heading
    # First extract title and subtitle
    title_match = re.match(r'^#\s+(.*?)\n##\s+(.*?)\n', md, re.MULTILINE)

    # ===== Cover Page =====
    for _ in range(4):
        doc.add_paragraph()

    if title_match:
        title_text = title_match.group(1).strip()
        subtitle_text = title_match.group(2).strip()
    else:
        title_text = '\u5fae\u4fe1\u652f\u4ed8\u62a5\u544a'
        subtitle_text = ''

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(title_text)
    set_font(run, 28, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

    if subtitle_text:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sp.add_run(subtitle_text)
        set_font(run, 16, color=RGBColor(0x66, 0x66, 0x66))

    add_hr(doc)

    # Extract metadata block (the first blockquote)
    meta_match = re.search(r'(?:^>\s+\*\*(.+?)\*\*[：:]\s*(.+)$)+', md, re.MULTILINE)
    meta_lines = re.findall(r'^>\s+\*\*(.+?)\*\*[：:]\s*(.+)$', md, re.MULTILINE)
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}\uff1a")
        set_font(run_l, 11, bold=True, color=RGBColor(0x33, 0x33, 0x33))
        run_v = p.add_run(value)
        set_font(run_v, 11, color=RGBColor(0x2B, 0x57, 0x9A))
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== Body =====
    # Remove title, subtitle, and metadata block; start from ## 目录
    body_start = md.find('\n## ')
    if body_start == -1:
        body_start = 0
    body_md = md[body_start:]

    process_md(body_md, doc)

    # ===== Footer =====
    doc.add_paragraph()
    add_hr(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('\u2014 \u62a5\u544a\u7ed3\u675f \u2014')
    set_font(run, 10, color=RGBColor(0x99, 0x99, 0x99))

    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp2.add_run('\u62a5\u544a\u751f\u6210\u65f6\u95f4\uff1a2026-03-22 | \u4ec5\u4f9b\u5185\u90e8\u6218\u7565\u7814\u7a76\u4f7f\u7528')
    set_font(run, 9, color=RGBColor(0xAA, 0xAA, 0xAA))

    # Save
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       '\u5fae\u4fe1\u652f\u4ed8\u7b2c\u4e00\u9009\u62e9\u6218\u7565\u7814\u7a76\u62a5\u544a.docx')
    doc.save(out)
    print(f'\u2705 Word\u62a5\u544a\u5df2\u751f\u6210\uff1a{out}')
    return out


if __name__ == '__main__':
    generate()
