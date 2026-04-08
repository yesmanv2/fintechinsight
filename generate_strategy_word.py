#!/usr/bin/env python3
"""生成微信支付第一选择战略研究报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A3C6E")
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "E8EDF3")
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(32)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(26)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x3D, 0x7E, 0xBB)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(22)
    return p

def add_para(doc, text, bold=False, italic=False, font_size=10.5, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_rich_para(doc, segments, space_after=6):
    """添加富文本段落，segments为列表：(text, bold, color)"""
    p = doc.add_paragraph()
    for seg in segments:
        if isinstance(seg, str):
            text, is_bold, clr = seg, False, None
        elif len(seg) == 2:
            text, is_bold = seg
            clr = None
        else:
            text, is_bold, clr = seg
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if is_bold:
            run.bold = True
        if clr:
            run.font.color.rgb = clr
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.name = '微软雅黑'
        run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_callout(doc, text, bg_color="E8F0FE", border_color="2B579A"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_numbered(doc, items):
    """添加编号列表"""
    for i, item in enumerate(items, 1):
        if isinstance(item, tuple):
            bold_part, rest = item
            p = doc.add_paragraph()
            run_num = p.add_run(f"{i}. ")
            run_num.font.size = Pt(10)
            run_num.font.name = '微软雅黑'
            run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run_b = p.add_run(bold_part)
            run_b.font.size = Pt(10)
            run_b.font.name = '微软雅黑'
            run_b._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run_b.bold = True
            run_t = p.add_run(rest)
            run_t.font.size = Pt(10)
            run_t.font.name = '微软雅黑'
            run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {item}")
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.left_indent = Cm(0.5)


def generate_report():
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.15
    
    # =============================================
    # 封面
    # =============================================
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = Pt(42)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("让微信支付成为用户的第一选择“)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.line_spacing = Pt(22)
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(”—— 危与机深度战略研究报告“)
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    add_hr(doc)
    
    info_items = [
        (”报告定位“, ”内部战略研究文档，仅供核心团队参考“),
        (”研究时间", "2026年3月“),
        (”数据来源“, ”小红书全平台舆情7162条 + 外部行业研究 + 黑猫投诉 + 公开市场数据"),
        ('研究目标', '围绕\u201c让微信支付成为用户的第一选择\u201d这一战略目标，寻找危与机'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}：“)
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        run_l.font.name = '微软雅黑'
        run_l._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_l.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.name = '微软雅黑'
        run_v._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_v.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # =============================================
    # 目录
    # =============================================
    add_heading_styled(doc, ”目  录“, level=1)
    add_hr(doc)
    
    toc_items = [
        (”一、核心结论“, False),
        (”二、微信支付可以改进的方向", False),
        ("    2.1 用户安全感知：风控精度与客服响应的提升空间", True),
        ("    2.2 金融产品矩阵：从“有”到“有感”的进阶空间", True),
        ("    2.3 自动续费管理：化行业痛点为品牌优势的契机", True),
        ("    2.4 海外支付：从“覆盖有”到“体验优”的升级空间", True),
        ("    2.5 商户端服务：从“标准化”到"差异化赋能"", True),
        ("三、五大竞品做对了什么", False),
        ("    3.1 支付宝：金融产品矩阵的"飞轮效应"", True),
        ("    3.2 云闪付：费率优势 + 银行信任", True),
        ("    3.3 抖音支付：内容电商闭环", True),
        ("    3.4 美团支付：垂直场景统治", True),
        ("    3.5 京东支付：电商分期场景“, True),
        (”四、微信支付的战略机会“, False),
        (”五、微信支付面临的风险“, False),
        (”六、用户为什么选择别人而不是我们“, False),
        (”七、战略行动路线图“, False),
        (”附录：关键数据一览“, False),
    ]
    for text, is_sub in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(text.strip())
        run.font.size = Pt(11 if not is_sub else 10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if not is_sub:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        else:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_after = Pt(4)
        if is_sub:
            p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_page_break()
    
    # =============================================
    # 一、核心结论
    # =============================================
    add_heading_styled(doc, ”一、核心结论“, level=1)
    add_hr(doc)
    
    add_heading_styled(doc, ”一句话判断“, level=3)
    add_callout(doc,
        ”微信支付正在从“默认选择”滑向“被动选择”。用户不是不用微信支付，而是越来越多场景下“不得不用”而非“想用”——这是最危险的信号。当竞品在各自场景中把“不得不用”变成“更想用”时，微信支付的市场份额将被逐步蚕食。",
        bg_color="FFF3E0", border_color="E65100“
    )
    
    add_heading_styled(doc, ”三个核心判断", level=3)
    create_styled_table(doc,
        ["#", "判断“, ”紧迫度"],
        [
            ["1“, ”微信支付的核心优势正在从“场景覆盖广”退化为“社交绑定“ — 当用户只在转账、发红包时想到微信支付，而在购物、出行、理财、借贷时首选别人，”第一选择”就只剩了半壁江山", "🔴 高"],
            ["2“, ”金融产品矩阵是最大短板 — 支付宝的“余额宝+花呗+芝麻信用”铁三角已形成完整用户价值闭环，微信支付的“零钱通+分付+微粒贷”远未形成同等心智", "🔴 高"],
            ["3“, ”场景垄断正在被打破 — 抖音支付在内容电商、美团支付在本地生活、云闪付在公共交通，各自建立了独立闭环，用户的支付习惯正在被分流", "🟡 中“],
        ],
        col_widths=[1, 12.5, 2.5]
    )
    
    doc.add_page_break()
    
    # =============================================
    # 二、微信支付可以改进的方向
    # =============================================
    add_heading_styled(doc, ”二、微信支付可以改进的方向“, level=1)
    add_hr(doc)
    add_callout(doc, ”微信支付已经做了很多正确的事，但在通往“用户第一选择”的路上，以下方向仍有提升空间。")
    
    # 2.1
    add_heading_styled(doc, "2.1 用户安全感知：风控精度与客服响应的提升空间“, level=2)
    add_para(doc, ”这是当前用户反馈中最集中的方向，也是提升潜力最大的。“, bold=True)
    add_para(doc, ”用户把钱放在微信里，对资金安全的感知直接影响信任度。目前在风控精准度和异常情况处理链路上，还有优化空间。“)
    
    add_para(doc, ”数据证据：", bold=True, space_after=4)
    add_bullet(doc, " 风控误伤/账户冻结投诉 40条，月环比增长 185.7%")
    add_bullet(doc, " 转账/安全问题投诉 45条，月环比增长 433.3%")
    add_bullet(doc, " 两个话题共现 21次（所有话题对中最高），说明是同一个系统性问题的两面")
    add_bullet(doc, " 客服体验差投诉 21条，平均互动量 7555.3（所有话题中最高）“)
    
    add_para(doc, ”改进空间：“, bold=True, space_after=4)
    add_para(doc, ”当前的风控策略在合规和反洗钱上做得很到位，但在后续用户体验链路上可以进一步优化：“)
    add_numbered(doc, [
        (”冻结时可以更主动地告知原因", " → 降低用户恐慌“),
        (”人工客服入口可以更显眼", " → 让用户在焦虑时快速找到帮助“),
        (”申诉流程可以进一步简化", " → 缩短用户等待和焦虑周期“),
        (”解冻后增加关怀机制", " → 修复用户信任“),
    ])
    
    add_para(doc, ”行业参考：支付宝的资金安全体系有“账户安全险”、“被盗全赔”承诺、以及相对通畅的人工客服通道，这些做法值得借鉴。", italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=10)
    
    add_hr(doc)
    
    # 2.2
    add_heading_styled(doc, "2.2 金融产品矩阵：从“有”到“有感”的进阶空间“, level=2)
    add_para(doc, ”微信支付的金融产品线已经齐全，但在用户心智占有率上还有很大提升空间。“, bold=True)
    
    create_styled_table(doc,
        [”产品“, ”微信支付“, ”支付宝对标“, ”提升方向“],
        [
            [”理财“, ”零钱通“, ”余额宝“, ”余额宝是“国民理财”代名词，零钱通的用户认知度还可以进一步提高“],
            [”信贷“, ”分付/微粒贷“, ”花呗/借呗“, ”花呗已是线下支付标配选项，分付在商户端覆盖可以进一步拓展“],
            [”信用“, ”无独立体系“, ”芝麻信用“, ”芝麻信用渗透租车、酒店、签证等300+场景，微信有机会建设对标体系“],
            [”分期“, ”分付分期“, ”花呗分期免息“, ”花呗在大促节点大规模推免息，分付的营销力度可以加大“],
        ],
        col_widths=[2, 3, 3, 8]
    )
    
    doc.add_paragraph()
    add_para(doc, ”提升方向：", bold=True, space_after=4)
    add_bullet(doc, " 零钱通的入口可以更显眼，像余额宝那样在首页展示收益")
    add_bullet(doc, " 分付的开通引导和使用场景可以更丰富")
    add_bullet(doc, " 微粒贷与微信支付的品牌关联可以进一步强化")
    add_bullet(doc, " 可以探索类似芝麻信用的“信用积分”体系，串联所有金融产品")
    
    add_hr(doc)
    
    # 2.3
    add_heading_styled(doc, "2.3 自动续费管理：化行业痛点为品牌优势的契机“, level=2)
    add_para(doc, ”自动续费是全行业共同面临的用户体验挑战，微信支付作为行业领导者，有机会率先树立标杆。“, bold=True)
    
    add_para(doc, ”数据证据：", bold=True, space_after=4)
    add_bullet(doc, " 自动续费/免密支付投诉 64条（所有话题中第一名）")
    add_bullet(doc, " 平均互动量 2435.5")
    add_bullet(doc, " 黑猫投诉2025年全年自动续费类投诉达 69,201件，三年累计超 22万条“)
    
    add_para(doc, ”可优化方向：“, bold=True, space_after=4)
    add_numbered(doc, [
        (”简化取消流程", " — 目前取消路径较长，可以缩短为2-3步“),
        (”强化续费前提醒", " — 在扣款前主动推送消息通知，让用户有充分的知情权“),
        (”平衡商家与用户利益", " — 在保障商家合理权益的同时，更好地保护用户的选择权“),
        (”主动做合规标杆", " — 行业规范正在趋严，率先落实会赢得用户好感和监管认可"),
    ])
    
    add_hr(doc)
    
    # 2.4
    add_heading_styled(doc, "2.4 海外支付：从“覆盖有”到“体验优”的升级空间“, level=2)
    add_para(doc, ”微信支付在海外支付方面已经有了布局，接下来的重点是从覆盖广度转向体验深度。", bold=True)
    add_bullet(doc, " 汇率透明度可以进一步提升")
    add_bullet(doc, " 境外商户覆盖密度与支付宝尚有差距")
    add_bullet(doc, " 外卡绑定流程偏复杂，港澳台用户反馈较多“)
    add_para(doc, ”数据证据：微信支付海外类笔记 211条，正面率仅 10.7%。支付宝在海外商户覆盖、跨境支付成功率、汇率透明度上目前领先。", italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=10)
    
    add_hr(doc)
    
    # 2.5
    add_heading_styled(doc, "2.5 商户端服务：从“标准化”到"差异化赋能"", level=2)
    add_para(doc, "微信支付在商户端建立了庞大的覆盖网络，下一步可以在中小商户的服务深度上进一步发力。", bold=True)
    add_bullet(doc, " 商户手续费率（0.6%为主）在行业中处于中等水平，可以探索更有竞争力的费率方案")
    add_bullet(doc, " 结算周期标准为T+1，部分竞品在特定场景可做到实时到账，可以探索提速")
    add_bullet(doc, " 商户端工具（数据分析、营销功能）可以更加丰富")
    add_bullet(doc, " 从“基础支付服务”升级为“商户经营赋能平台”，增强商户粘性“)
    
    doc.add_page_break()
    
    # =============================================
    # 三、五大竞品做对了什么
    # =============================================
    add_heading_styled(doc, ”三、五大竞品做对了什么“, level=1)
    add_hr(doc)
    add_callout(doc, ”本章只讨论竞品做对了而微信支付没做到或没做好的方面，供内部参考借鉴。")
    
    # 3.1 支付宝
    add_heading_styled(doc, "3.1 支付宝：金融产品矩阵的"飞轮效应"", level=2)
    add_para(doc, "支付宝做对了最重要的一件事：把“支付工具”做成了“金融生活平台”。“, bold=True)
    
    add_para(doc, ”飞轮模型：支付 → 余额宝(理财) → 芝麻信用(信用评分) → 花呗(信贷) → 更多支付 → 用户粘性闭环“, italic=True, font_size=10)
    
    create_styled_table(doc,
        [”维度“, ”支付宝做了什么“, ”为什么有效“],
        [
            [”余额宝", "2013年推出至今，“把钱放余额宝”已成为国民习惯“, ”让用户的钱“停在”支付宝体系内“],
            [”芝麻信用“, ”渗透免押金租车/酒店/共享单车等300+场景“, ”信用体系创造了不可替代性“],
            [”花呗“, ”线上+线下全覆盖，商户端主动推荐花呗分期“, ”让用户形成“先买后付”的习惯“],
            [”碰一下", "2025年推出NFC碰付，线下体验比扫码更快“, ”直接竞争微信支付的线下扫码场景“],
        ],
        col_widths=[2.5, 6.5, 7]
    )
    
    doc.add_paragraph()
    add_para(doc, ”数据佐证：", bold=True, space_after=4)
    add_bullet(doc, " 支付宝正面笔记 222条（6平台第一），TOP1优势维度是“体验好”（53条）")
    add_bullet(doc, " 支付宝在“安全感”维度获得 13条正面评价，微信支付仅 0条")
    
    add_hr(doc)
    
    # 3.2 云闪付
    add_heading_styled(doc, "3.2 云闪付：费率优势 + 银行信任“, level=2)
    add_para(doc, ”云闪付在微信支付和支付宝的夹缝中找到了“低费率+银行安全感”的差异化定位。“, bold=True)
    
    create_styled_table(doc,
        [”维度“, ”云闪付做了什么“, ”效果“],
        [
            [”费率“, ”大量"0费率“”优惠活动"", "正面笔记中“费率低”占54条（49%）“],
            [”银行直连“, ”直接连接银行账户，不需要“零钱”中间层“, ”用户觉得钱直接从银行出，更安全“],
            [”公共交通“, ”深耕公交、地铁、高速收费等场景“, ”成为很多城市公共交通的首选支付方式“],
            [”银行活动“, ”各银行在云闪付上投放大量优惠活动“, ”用户养成了“有优惠就用云闪付”的习惯“],
        ],
        col_widths=[2.5, 6.5, 7]
    )
    
    doc.add_paragraph()
    add_callout(doc,
        ”注：云闪付在本次数据中笔记总量较少（596条，为6平台最少），相关结论需结合更大样本量验证。其费率优势和银行直连的差异化定位值得持续关注。",
        bg_color="FFF8E1", border_color="F9A825"
    )
    
    add_hr(doc)
    
    # 3.3 抖音支付
    add_heading_styled(doc, "3.3 抖音支付：内容电商闭环 + 本地生活分流“, level=2)
    add_para(doc, ”抖音支付做对了最重要的一件事：把支付嵌入内容消费的天然链路中，用户“看着看着就付了钱”。“, bold=True)
    
    create_styled_table(doc,
        [”维度“, ”抖音支付做了什么“, ”效果“],
        [
            [”内容-支付闭环“, ”直播/短视频→商品→抖音支付，全链路不出APP", "2025年支付GMV超 8500亿（+59%）“],
            [”商家扶持“, ”大规模免佣+低费率，吸引商家接入抖音支付“, ”商家有动力引导用户使用抖音支付“],
            [”月付(DOU分期)", "类似花呗的先享后付产品“, ”减少了用户在抖音内选择微信支付的动机“],
            [”本地生活“, ”团购、到店消费场景深度绑定“, ”直接与微信支付的线下场景竞争“],
        ],
        col_widths=[3, 6.5, 6.5]
    )
    
    doc.add_paragraph()
    add_para(doc, ”影响分析：抖音正在建立独立于微信的完整支付生态——以前用户在抖音买东西会用微信支付，现在越来越多用抖音自己的支付。2026年目标新增4000亿+GMV。", font_size=10)
    
    add_hr(doc)
    
    # 3.4 美团支付
    add_heading_styled(doc, "3.4 美团支付：垂直场景的绝对统治“, level=2)
    add_para(doc, ”美团支付在餐饮外卖这个高频场景中建立了“无感支付”体验。“, bold=True)
    
    create_styled_table(doc,
        [”维度“, ”美团支付做了什么“, ”效果“],
        [
            [”场景锁定“, ”外卖、到店、酒旅全场景默认美团支付“, ”用户点外卖时很少会主动切换到微信支付“],
            [”美团月付“, ”免息期+低门槛开通“, ”在高频消费场景中培养了“先吃后付”的习惯“],
            [”优惠券绑定“, ”红包、优惠券与美团支付绑定使用“, ”用美团支付比用微信支付多省几元钱“],
            [”会员体系“, ”美团会员权益与美团支付深度绑定“, ”增加了用户切换成本“],
        ],
        col_widths=[2.5, 6.5, 7]
    )
    
    doc.add_paragraph()
    add_para(doc, ”影响分析：美团正面笔记 140条，在餐饮外卖这个日频次最高的支付场景中，美团支付已经成为默认选项。", font_size=10)
    
    add_hr(doc)
    
    # 3.5 京东支付
    add_heading_styled(doc, "3.5 京东支付：电商场景 + 白条信贷“, level=2)
    add_para(doc, ”京东支付通过白条产品在电商场景中建立了深度绑定。“, bold=True)
    
    create_styled_table(doc,
        [”维度“, ”京东支付做了什么“, ”效果“],
        [
            [”京东白条“, ”购物分期+免息活动，在京东购物时默认推荐“, ”用户在京东购物形成“白条付”习惯“],
            [”费率优势“, ”正面笔记中“费率低”占13条（28%）“, ”在有限的正面评价中，费率是第一优势“],
            [”先享后付“, ”可以先收货确认再付款“, ”降低了用户的支付心理障碍“],
        ],
        col_widths=[2.5, 6.5, 7]
    )
    
    add_para(doc, ”影响：相对有限，京东支付的场景比较垂直。但京东白条在大额消费分期场景中的竞争力不容忽视。“, font_size=10, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    
    doc.add_page_break()
    
    # =============================================
    # 四、微信支付的战略机会
    # =============================================
    add_heading_styled(doc, ”四、微信支付的战略机会“, level=1)
    add_hr(doc)
    add_callout(doc, ”基于竞品分析和用户洞察，微信支付有以下未被充分利用的机会。")
    
    # 4.1
    add_heading_styled(doc, "4.1 🔵 最大机会：信任资产重建“, level=2)
    add_para(doc, ”在所有支付平台中，用户对“资金安全”的需求是最刚性的。微信支付在安全感知上还有提升空间，而微信的社交属性意味着微信支付比任何竞品都更“近”用户——如果能在这个距离上建立安全感，将是最强的竞争壁垒。“, bold=True)
    
    add_para(doc, ”具体方向：“, bold=True, space_after=4)
    add_numbered(doc, [
        ”推出“微信支付安心保障计划”——账户被盗全赔、误操作72小时退款“,
        ”风控冻结时主动通知原因+预计解冻时间+一键申诉入口“,
        ”在钱包首页增加“资金安全中心”入口，让用户可视化自己的安全保障状态",
    ])
    
    # 4.2
    add_heading_styled(doc, "4.2 🔵 第二大机会：金融产品矩阵升级“, level=2)
    add_para(doc, ”微信支付的零钱通、分付、微粒贷在产品层面并不差，但用户心智占有率极低。这既是问题也是机会。“, bold=True)
    
    add_para(doc, ”具体方向：“, bold=True, space_after=4)
    add_numbered(doc, [
        ”零钱通升级：在微信支付主页显眼位置展示"今日收益"",
        "分付破局：联合小程序电商做大规模“分付免息”活动“,
        ”信用体系建设：基于微信实名认证+社交图谱+支付历史，构建"微信信用分"",
        "理财频道整合：在微信支付内建立“理财”入口",
    ])
    
    # 4.3
    add_heading_styled(doc, "4.3 🔵 第三大机会：自动续费合规先手“, level=2)
    add_para(doc, ”自动续费是全行业第一大投诉话题，但也是最好的品牌差异化机会——谁先真正解决这个问题，谁就赢得用户信任。“, bold=True)
    add_numbered(doc, [
        ”推出“续费提醒助手”——每次续费前5天推送微信消息通知，一键取消“,
        ”在钱包中增加“订阅管理中心”——统一展示所有自动续费服务，一键全部关闭“,
        ”公开发布“微信支付自动续费管理白皮书”，倡导行业规范",
    ])
    
    # 4.4
    add_heading_styled(doc, "4.4 🔵 第四大机会：社交支付的不可替代性深化“, level=2)
    add_para(doc, ”微信红包、转账、群收款是其他任何平台都无法复制的场景。可以从“工具”升级为“金融关系”。“, bold=True)
    add_numbered(doc, [
        ”亲情账户：父母/子女之间设立共享理财账户，支持代付和额度管理“,
        ”社交化理财：好友之间的“储蓄挑战“”理财心愿单”等社交化产品“,
        ”红包升级：把红包从“钱”升级为“权益”——可以发优惠券红包、会员权益红包等",
    ])
    
    # 4.5
    add_heading_styled(doc, "4.5 🔵 第五大机会：AI + 支付的融合创新", level=2)
    add_para(doc, "2026年AI应用爆发，支付是AI最天然的商业化场景之一。“, bold=True)
    add_numbered(doc, [
        ”智能账单助手：基于AI自动分类账单、分析消费习惯、推荐省钱方案",
        "AI风控升级：用大模型优化风控准确率，降低误伤率“,
        ”智能客服重塑：AI客服能真正解决问题，而不是"踢皮球"",
        "刷掌支付规模化：在线下场景推广生物识别支付，建立技术壁垒“,
    ])
    
    doc.add_page_break()
    
    # =============================================
    # 五、微信支付面临的风险
    # =============================================
    add_heading_styled(doc, ”五、微信支付面临的风险", level=1)
    add_hr(doc)
    
    # 5.1
    add_heading_styled(doc, "5.1 🔴 最大风险：场景分流“, level=2)
    add_para(doc, ”每一个超级APP都在构建自己的支付闭环。当用户在抖音买东西用抖音支付、在美团点餐用美团支付、在京东购物用京东白条，微信支付的使用场景将被逐步挤压。“, bold=True)
    
    create_styled_table(doc,
        [”平台“, ”分流的场景“, ”分流的支付GMV量级“, ”趋势“],
        [
            [”抖音支付“, ”直播电商、短视频购物、本地生活团购", "8500亿+/年，年增59%", "🔺 加速增长“],
            [”美团支付“, ”外卖、到店餐饮、酒旅“, ”数千亿级（未公开）", "🔺 稳步增长“],
            [”云闪付“, ”公共交通、政务缴费、线下消费“, ”千亿级", "🔺 持续增长“],
            [”京东支付“, ”电商购物（尤其3C/家电大额分期）“, ”千亿级", "➡️ 平稳“],
            [”支付宝“, ”线上购物、理财、信贷、出行“, ”万亿级", "➡️ 稳固“],
        ],
        col_widths=[2.5, 5.5, 4.5, 3.5]
    )
    
    doc.add_paragraph()
    add_callout(doc, 
        ”叠加效应：一个用户可能——早上用美团支付点外卖、中午用抖音支付买东西、下午坐地铁用银联、晚上用支付宝花呗分期购物——一天之中，微信支付的使用次数可能为零。",
        bg_color="FFF3E0", border_color="E65100"
    )
    
    # 5.2
    add_heading_styled(doc, "5.2 🔴 第二大风险：信任赤字累积“, level=2)
    add_para(doc, ”风控误伤、客服差、自动续费等问题叠加，正在累积“信任赤字”。")
    add_bullet(doc, " 微信支付的强流失信号（15条）在6大平台中排名第一")
    add_bullet(doc, " 负面率趋势“连续上升”——是所有平台中唯一被标记为持续上升的")
    add_bullet(doc, " 极度愤怒+强烈不满用户占负面帖子的 18%")
    
    # 5.3
    add_heading_styled(doc, "5.3 🟡 第三大风险：数字人民币的潜在冲击“, level=2)
    add_para(doc, ”央行推进的数字人民币（e-CNY）可能在2026-2027年大规模推广。作为“法定货币的数字形态”，数字人民币不需要绑定银行卡、没有手续费——如果用户习惯了直接用数字人民币，微信支付将面临“去中介化”风险。")
    
    # 5.4
    add_heading_styled(doc, "5.4 🟡 第四大风险：监管趋严下的合规成本“, level=2)
    create_styled_table(doc,
        [”监管方向“, ”对微信支付的潜在影响“],
        [
            [”统一风控模型“, ”可能需要调整现有风控策略，短期内误伤率可能波动“],
            [”穿透式监管“, ”交易数据实时报送要求提高，系统成本增加“],
            [”反洗钱升级“, ”实名认证门槛可能提高，新用户开通体验可能变差“],
            [”现金支付保障“, ”线下商户必须接受现金，移动支付的“独占场景”减少“],
        ],
        col_widths=[4, 12]
    )
    
    doc.add_page_break()
    
    # =============================================
    # 六、用户为什么选择别人而不是我们
    # =============================================
    add_heading_styled(doc, ”六、用户为什么选择别人而不是我们", level=1)
    add_hr(doc)
    
    add_heading_styled(doc, "6.1 用户支付决策的四层模型", level=2)
    add_callout(doc,
        '第一层：默认选择 → \u201c不用想，手自然就扫了\u201d（习惯驱动）\n'
        '第二层：利益选择 → \u201c这个更便宜/有优惠\u201d（优惠驱动）\n'
        '第三层：功能选择 → \u201c只有这个能做到\u201d（功能驱动）\n'
        '第四层：信任选择 → \u201c大钱放在更安全的地方\u201d（安全驱动）'
    )
    
    add_heading_styled(doc, "6.2 微信支付在每一层的竞争地位“, level=2)
    create_styled_table(doc,
        [”决策层“, ”微信支付的位置“, ”对比“, ”分析“],
        [
            [”第一层：默认选择", "🟡 部分场景仍是默认“, ”美团(外卖)、抖音(直播购物)", "这些APP内的默认支付已经不是微信支付“],
            [”第二层：利益选择", "🟡 有提升空间“, ”部分竞品有更多优惠活动“, ”微信支付可以加大优惠活动的力度和频次“],
            [”第三层：功能选择", "🟡 有提升空间“, ”支付宝(芝麻信用、余额宝)", "微信支付的金融功能可以更突出“],
            [”第四层：信任选择", "🟡 中等“, ”支付宝(保险)、银行直连类支付“, ”微信支付可以加强安全保障的可感知度"],
        ],
        col_widths=[3, 3, 4.5, 5.5]
    )
    
    add_heading_styled(doc, "6.3 五个典型用户流失场景“, level=2)
    
    scenarios = [
        ('场景一：购物分期', '用户想买一台5000元的手机。支付宝弹出\u201c花呗12期免息\u201d，微信支付没有同等推送。\u2192 用户选择支付宝。'),
        ('场景二：日常理财', '用户发了工资，想找个地方放。朋友推荐\u201c余额宝，每天能看到收益\u201d。\u2192 用户把钱转到余额宝。'),
        ('场景三：线下坐车', '用户到了地铁站，发现有\u201c1分钱坐地铁\u201d优惠活动。\u2192 用户出行场景的默认支付习惯被改变。'),
        ('场景四：刷抖音购物', '用户看直播时被种草了一件衣服，点了\u201c立即购买\u201d。抖音弹出\u201c用月付，免息3期\u201d。\u2192 用户用了抖音支付。'),
        ('场景五：外卖消费', '用户点外卖，美团弹出\u201c用美团支付再减3元\u201d。\u2192 每天一单外卖，美团支付占据了用户最高频的支付场景。'),
    ]
    for title_text, desc in scenarios:
        add_para(doc, title_text, bold=True, font_size=10.5, space_after=2)
        add_callout(doc, desc)
    
    add_para(doc, '共同规律：每一个场景流失都遵循\u201c优惠/便利触发 \u2192 首次使用 \u2192 习惯养成 \u2192 默认绑定\u201d的路径。微信支付不是在某一刻失去了用户，而是在每一个场景中逐渐不再被需要。', bold=True)
    
    doc.add_page_break()
    
    # =============================================
    # 七、战略行动路线图
    # =============================================
    add_heading_styled(doc, ”七、战略行动路线图“, level=1)
    add_hr(doc)
    
    add_heading_styled(doc, '7.1 总体战略：”三守一攻“', level=2)
    create_styled_table(doc,
        [”战略“, ”含义“, ”时间框架“],
        [
            [”守信任“, ”重建资金安全感，消除信任赤字", "0-3个月（紧急）“],
            [”守体验“, ”解决自动续费、客服、风控误伤三大痛点", "0-6个月（重要）“],
            [”守场景“, ”巩固社交支付+小程序生态的核心壁垒“, ”持续“],
            [”攻金融“, ”金融产品矩阵升级，从“支付工具”进化为"金融生活平台"", "6-18个月（战略）"],
        ],
        col_widths=[2.5, 8, 5.5]
    )
    
    add_heading_styled(doc, "7.2 紧急行动（0-3个月）“, level=2)
    
    add_para(doc, ”行动1：信任修复计划“, bold=True, font_size=11)
    create_styled_table(doc,
        [”动作“, ”具体措施“, ”预期效果“],
        [
            [”推出"安心保障"", "账户被盗全赔承诺 + 可视化保障页面“, ”提升用户安全感知“],
            [”风控透明化“, ”冻结时主动通知原因+预计时间+一键申诉“, ”降低风控误伤投诉50%+“],
            [”客服升级“, ”增加人工客服入口、赋予更多权限、承诺响应时效“, ”降低客服投诉“],
        ],
        col_widths=[3.5, 7, 5.5]
    )
    
    doc.add_paragraph()
    add_para(doc, ”行动2：自动续费治理“, bold=True, font_size=11)
    create_styled_table(doc,
        [”动作“, ”具体措施“, ”预期效果“],
        [
            [”续费提醒“, ”扣费前5天推送微信消息，含一键取消按钮“, ”降低自动续费投诉60%+“],
            [”订阅管理中心“, ”钱包内统一展示所有自动续费服务“, ”提升用户掌控感“],
            [”行业标准倡导“, ”发布白皮书，推动行业规范“, ”品牌美誉度提升"],
        ],
        col_widths=[3.5, 7, 5.5]
    )
    
    add_heading_styled(doc, "7.3 重要行动（3-6个月）“, level=2)
    
    add_para(doc, ”行动3：金融产品可见性提升“, bold=True, font_size=11)
    create_styled_table(doc,
        [”产品“, ”具体措施“, ”目标“],
        [
            [”零钱通“, ”首页展示“今日收益”、推送月度收益报告“, ”认知率提升至50%+“],
            [”分付“, ”联合小程序商家做免息活动、在支付确认页默认推荐“, ”开通率提升3倍“],
            [”理财入口“, ”在微信支付主页增加"理财"Tab", "日均UV提升10倍“],
        ],
        col_widths=[3, 8.5, 4.5]
    )
    
    doc.add_paragraph()
    add_para(doc, ”行动4：海外支付体验升级“, bold=True, font_size=11)
    create_styled_table(doc,
        [”动作“, ”具体措施“, ”目标“],
        [
            [”汇率透明“, ”支付时展示实时汇率和银行对比“, ”消除汇率吐槽“],
            [”商户覆盖“, ”加速境外商户签约“, ”追平支付宝覆盖率“],
            [”外卡体验“, ”简化港澳台及外卡绑定流程“, ”降低绑卡失败率"],
        ],
        col_widths=[3, 8.5, 4.5]
    )
    
    add_heading_styled(doc, "7.4 战略行动（6-18个月）“, level=2)
    
    add_para(doc, ”行动5：微信信用体系建设“, bold=True, font_size=11)
    create_styled_table(doc,
        [”阶段“, ”具体措施“, ”目标“],
        [
            [”第一期“, ”基于支付数据+实名信息构建基础信用评分“, ”体系上线“],
            [”第二期“, ”免押金接入共享单车、充电宝等高频场景“, ”用户感知信用价值“],
            [”第三期“, ”拓展到租房、酒店、租车等更多场景“, ”构建迁移成本壁垒“],
        ],
        col_widths=[2.5, 8.5, 5]
    )
    
    doc.add_paragraph()
    add_para(doc, ”行动6：社交支付差异化深化“, bold=True, font_size=11)
    create_styled_table(doc,
        [”产品“, ”具体措施“, ”目标“],
        [
            [”亲情账户“, ”共享理财+代付+额度管理“, ”锁定家庭支付场景“],
            [”社交理财“, ”储蓄挑战、理财心愿单“, ”利用社交关系促活“],
            [”权益红包“, ”优惠券红包、会员权益红包“, ”红包从“工具”升级为"平台""],
        ],
        col_widths=[2.5, 8.5, 5]
    )
    
    doc.add_paragraph()
    add_para(doc, "行动7：AI + 支付融合“, bold=True, font_size=11)
    create_styled_table(doc,
        [”产品“, ”具体措施“, ”目标“],
        [
            [”智能账单", "AI自动分类+消费分析+省钱建议“, ”提升金融工具价值感"],
            ["AI客服“, ”大模型驱动的客服系统，能真正解决问题“, ”客服满意度翻倍"],
            ["AI风控“, ”优化风控模型准确率“, ”误伤率降低80%+“],
        ],
        col_widths=[2.5, 8.5, 5]
    )
    
    doc.add_page_break()
    
    # =============================================
    # 附录
    # =============================================
    add_heading_styled(doc, ”附录：关键数据一览", level=1)
    add_hr(doc)
    
    add_heading_styled(doc, "A. 六大平台舆情数据概览“, level=2)
    create_styled_table(doc,
        [”平台“, ”笔记总数“, ”正面率“, ”负面率“, ”客诉占比“, ”强流失信号“],
        [
            [”微信支付", "1,683", "10.7%", "11.8%", "3.3%", "15条“],
            [”支付宝", "1,949", "11.4%", "17.1%", "0.0%", "11条“],
            [”抖音支付", "1,063", "8.7%", "18.9%", "6.5%", "6条“],
            [”美团支付", "999", "14.0%", "23.5%", "17.1%", "2条“],
            [”京东支付", "872", "5.3%", "12.8%", "0.0%", "12条“],
            [”云闪付", "596", "18.5%", "14.3%", "0.0%", "3条"],
        ],
        col_widths=[2.5, 2.5, 2, 2, 2, 2.5]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "B. 微信支付痛点热力图“, level=2)
    create_styled_table(doc,
        [”痛点“, ”投诉数“, ”平均互动量“, ”月环比“, ”严重度“],
        [
            [”自动续费/免密支付", "64", "2,435.5", "+36.4%", "🔴"],
            ["转账/安全问题", "45", "4,852.7", "+433.3%", "🔴"],
            ["风控误伤/账户冻结", "40", "3,527.7", "+185.7%", "🔴"],
            ["分付/借贷", "21", "4,896.1", "+100.0%", "🟡"],
            ["客服体验差", "21", "7,555.3", "+200.0%", "🔴"],
            ["提现/手续费", "8", "3,165.6", "+50.0%", "🟡"],
        ],
        col_widths=[4, 2, 3, 2.5, 2]
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, "C. 竞品关注度排名“, level=2)
    create_styled_table(doc,
        [”排名“, ”竞品“, ”关注理由“, ”样本量"],
        [
            ["1“, ”支付宝“, ”金融产品矩阵+信用体系，用户粘性强", "1,949条"],
            ["2“, ”抖音支付", "GMV高速增长+场景闭环，分流效应明显", "1,063条"],
            ["3“, ”美团支付“, ”高频场景默认绑定，外卖场景统治力强", "999条"],
            ["4“, ”京东支付“, ”电商分期场景锁定", "872条"],
            ["5“, ”云闪付“, ”费率优势+银行信任（样本量较少，需持续观察）", "596条“],
        ],
        col_widths=[1.5, 2.5, 8, 2.5]
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 结语
    add_hr(doc)
    add_callout(doc,
        ”结语：微信支付成为用户第一选择的关键不在于“做更多”，而在于“把关键的事做到极致”。“
        ”当前最有价值的切入点是提升用户的安全感知和信任——让用户在遇到问题时能快速得到帮助，“
        ”这种体验会比100次顺畅支付更能赢得口碑。先让用户放心，再让用户满意，最后让用户离不开“
        ”——这是通往“第一选择”的清晰路径。微信支付拥有最大的用户基数和最强的社交基因，“
        ”只要在关键体验上持续打磨，成为用户第一选择的目标完全可以实现。",
        bg_color="E8F5E9", border_color="1B5E20“
    )
    
    doc.add_paragraph()
    
    # 页脚
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_line.add_run(”— 报告结束 —“)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_info.add_run(”报告生成时间：2026-03-22 | 仅供内部战略研究使用“)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 保存
    output_path = os.path.join(os.path.dirname(__file__), ”微信支付第一选择战略研究报告.docx")
    doc.save(output_path)
    print(f"✅ Word报告已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    generate_report()
