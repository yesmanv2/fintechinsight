#!/usr/bin/env python3
"""
支付宝服务商动态监控报告生成器
从 alipay_merchant_data.json 读取变化记录，生成独立 Word 报告。

Usage:
    python3 generate_alipay_merchant_report.py
    python3 generate_alipay_merchant_report.py --year 2026 --month 4
    python3 generate_alipay_merchant_report.py --output-dir ./reports
"""

import os
import sys
import json
import argparse
from datetime import datetime, date
from collections import Counter, defaultdict

from generate_word_report import (
    set_cell_shading,
    create_styled_table,
    add_heading_styled,
    add_paragraph_styled,
    add_bullet_point,
    add_callout_box,
    add_horizontal_line,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'alipay_merchant_data.json')

BRAND_COLOR = RGBColor(0x16, 0x77, 0xFF)  # 支付宝蓝
BRAND_HEX = "1677FF"


def load_data():
    if not os.path.exists(DATA_FILE):
        print(f"❌ 数据文件不存在: {DATA_FILE}")
        print("💡 请先运行: python3 crawl_alipay_merchant.py")
        sys.exit(1)
    
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        raw = json.load(f)
    
    return raw.get('changes', []), raw.get('meta', {})


def generate_report(year=None, month=None, output_dir=None):
    """生成支付宝服务商动态监控报告"""
    changes, meta = load_data()
    
    today = date.today()
    report_year = year or today.year
    report_month = month or today.month
    report_date_str = today.strftime('%Y年%m月%d日').replace('年0', '年').replace('月0', '月')
    
    # 数据分析
    total_changes = len(changes)
    
    # 按页面分组
    page_changes = defaultdict(list)
    for c in changes:
        page_changes[c.get('page_name', '未知')].append(c)
    
    # 按变化类型统计
    type_stats = Counter(c.get('change_type', 'unknown') for c in changes)
    
    # 按月份统计
    monthly_stats = Counter()
    for c in changes:
        snap_date = c.get('snapshot_date', '')[:7]
        if snap_date:
            monthly_stats[snap_date] += 1
    
    # 本月变化
    month_prefix = f"{report_year}-{report_month:02d}"
    this_month_changes = [c for c in changes if c.get('snapshot_date', '').startswith(month_prefix)]
    
    type_labels = {
        'added': '新增', 'removed': '删除', 'modified': '修改', 'initial': '首次快照'
    }
    type_icons = {
        'added': '🆕', 'removed': '🗑️', 'modified': '✏️', 'initial': '📋'
    }
    
    # ==================== 创建文档 ====================
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = Pt(42)
    run = title_p.add_run("支付宝服务商动态监控报告")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("服务商后台页面变化监控与分析")
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    add_horizontal_line(doc)
    
    info_items = [
        ("数据来源", "支付宝服务商后台网页爬取"),
        ("报告时间", report_date_str),
        ("监控页面", f"{len(page_changes)} 个"),
        ("累计变化", f"{total_changes} 条"),
        ("本月变化", f"{len(this_month_changes)} 条"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}：")
        run_l.font.size = Pt(11)
        run_l.font.bold = True
        run_l.font.name = '微软雅黑'
        run_l._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.name = '微软雅黑'
        run_v._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_v.font.color.rgb = BRAND_COLOR
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ==================== 一、概览 ====================
    add_heading_styled(doc, "一、监控概览", level=1)
    add_horizontal_line(doc)
    
    if total_changes > 0:
        added = type_stats.get('added', 0)
        removed = type_stats.get('removed', 0)
        modified = type_stats.get('modified', 0)
        add_callout_box(doc,
            f"截至{report_date_str}，累计检测到 {total_changes} 次页面变化，"
            f"其中新增 {added} 次、修改 {modified} 次、删除 {removed} 次。"
            f"本月新增变化 {len(this_month_changes)} 条。",
            bg_color="E8F0FE", border_color=BRAND_HEX)
    else:
        add_callout_box(doc,
            "尚未检测到任何页面变化。如果是首次运行，这属于正常情况。"
            "后续监控将与本次快照对比，检测新的变化。",
            bg_color="E8F0FE", border_color=BRAND_HEX)
    
    doc.add_paragraph()
    
    # 变化类型统计
    if type_stats:
        add_paragraph_styled(doc, "变化类型统计：", bold=True, font_size=11, space_after=8)
        type_rows = []
        for t, cnt in type_stats.most_common():
            icon = type_icons.get(t, '📋')
            label = type_labels.get(t, t)
            pct = f"{cnt / total_changes * 100:.1f}%" if total_changes > 0 else "0%"
            type_rows.append([f"{icon} {label}", str(cnt), pct])
        create_styled_table(doc, ["变化类型", "次数", "占比"], type_rows, col_widths=[5, 3, 3])
    
    doc.add_page_break()
    
    # ==================== 二、各页面变化详情 ====================
    add_heading_styled(doc, "二、各页面变化详情", level=1)
    add_horizontal_line(doc)
    
    for page_idx, (page_name, page_chs) in enumerate(sorted(page_changes.items())):
        add_heading_styled(doc, f"2.{page_idx+1} {page_name}（{len(page_chs)} 次变化）", level=2)
        
        # 按时间倒序
        sorted_chs = sorted(page_chs, key=lambda x: x.get('detected_at', ''), reverse=True)
        
        detail_rows = []
        for c in sorted_chs[:15]:  # 最多显示15条
            ct = c.get('change_type', '')
            icon = type_icons.get(ct, '📋')
            label = type_labels.get(ct, ct)
            summary = c.get('change_summary', '')[:60]
            snap_date = c.get('snapshot_date', '')
            detail_rows.append([snap_date, f"{icon} {label}", summary])
        
        if detail_rows:
            create_styled_table(doc,
                ["日期", "类型", "变化摘要"],
                detail_rows,
                col_widths=[2.5, 2.5, 11])
        
        # 显示最新变化的详细内容
        if sorted_chs and sorted_chs[0].get('new_content'):
            doc.add_paragraph()
            latest = sorted_chs[0]
            if latest.get('new_content'):
                content_preview = latest['new_content'][:300]
                add_callout_box(doc,
                    f"最新变化内容预览：\n{content_preview}",
                    bg_color="F5F5F5", border_color="9E9E9E")
        
        doc.add_paragraph()
        
        if page_idx < len(page_changes) - 1:
            add_horizontal_line(doc)
    
    doc.add_page_break()
    
    # ==================== 三、月度趋势 ====================
    if monthly_stats:
        add_heading_styled(doc, "三、月度变化趋势", level=1)
        add_horizontal_line(doc)
        
        sorted_months = sorted(monthly_stats.keys())
        max_count = max(monthly_stats.values()) if monthly_stats else 1
        trend_rows = []
        for mk in sorted_months:
            cnt = monthly_stats[mk]
            parts = mk.split('-')
            label = f"{parts[0]}年{int(parts[1])}月" if len(parts) == 2 else mk
            bar_len = max(1, int(cnt / max_count * 30))
            trend_rows.append([label, str(cnt), "█" * bar_len])
        create_styled_table(doc, ["月份", "变化数", "趋势"], trend_rows, col_widths=[3, 2, 11])
        
        doc.add_page_break()
    
    # ==================== 四、本月变化清单 ====================
    if this_month_changes:
        add_heading_styled(doc, "四、本月变化清单", level=1)
        add_horizontal_line(doc)
        
        month_rows = []
        for c in sorted(this_month_changes, key=lambda x: x.get('detected_at', '')):
            page = c.get('page_name', '')
            ct = type_labels.get(c.get('change_type', ''), '未知')
            summary = c.get('change_summary', '')[:50]
            snap_date = c.get('snapshot_date', '')
            month_rows.append([snap_date, page, ct, summary])
        
        create_styled_table(doc,
            ["日期", "页面", "类型", "摘要"],
            month_rows,
            col_widths=[2.5, 3, 2, 8.5])
        
        doc.add_page_break()
    
    # ==================== 结论 ====================
    add_heading_styled(doc, "结论与关注点", level=1)
    add_horizontal_line(doc)
    
    if total_changes > 0:
        most_active_page = max(page_changes.items(), key=lambda x: len(x[1]))
        add_paragraph_styled(doc,
            f"变化最频繁的页面为「{most_active_page[0]}」（{len(most_active_page[1])} 次变化），建议重点关注。",
            bold=True, font_size=11)
    
    doc.add_paragraph()
    
    add_callout_box(doc,
        "建议每月运行一次监控，及时掌握支付宝服务商后台的政策、费率、产品能力变化。"
        "对于新增或修改类变化，建议第一时间评估对业务的影响。",
        bg_color="E8F5E9", border_color="1B5E20")
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— 报告结束 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"报告生成时间：{today.strftime('%Y-%m-%d')} | 数据源：支付宝服务商后台监控")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 保存
    if output_dir is None:
        output_dir = BASE_DIR
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"支付宝服务商动态报告_{report_year}年{report_month:02d}月.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"✅ 报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="支付宝服务商动态监控报告生成器")
    parser.add_argument('--year', type=int, default=None)
    parser.add_argument('--month', type=int, default=None)
    parser.add_argument('--output-dir', type=str, default=None)
    args = parser.parse_args()
    
    generate_report(year=args.year, month=args.month, output_dir=args.output_dir)
