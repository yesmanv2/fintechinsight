#!/usr/bin/env python3
"""
通用多平台客诉分析报告生成器
从 real_data.json 按平台过滤数据，动态生成 Word 文档报告。
支持：支付宝、抖音支付、美团支付、京东支付、云闪付

Usage:
    python3 generate_platform_report.py --platform 支付宝
    python3 generate_platform_report.py --all
    python3 generate_platform_report.py --platform 京东支付 --year 2026 --month 3
"""

import os
import sys
import json
import argparse
from datetime import datetime, date
from collections import Counter, defaultdict

# 复用 generate_word_report.py 的样式工具函数
from generate_word_report import (
    set_cell_shading,
    create_styled_table,
    add_heading_styled,
    add_paragraph_styled,
    add_bullet_point,
    add_callout_box,
    add_horizontal_line,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 平台配置
PLATFORM_CONFIG = {
    "支付宝": {
        "icon": "💙", "color": "#1677FF",
        "brand_rgb": RGBColor(0x16, 0x77, 0xFF),
        "brand_hex": "1677FF",
        "sub_products": ["花呗", "借呗", "余额宝", "芝麻信用", "碰一下", "蚂蚁森林", "蚂蚁财富", "好医保", "备用金", "网商贷"],
        "relevance_keywords": ["支付宝", "花呗", "借呗", "余额宝", "芝麻", "蚂蚁", "Alipay", "碰一下",
                               "网商银行", "蚂蚁森林", "蚂蚁财富", "冻结", "盗刷", "乱扣", "续费"],
    },
    "抖音支付": {
        "icon": "🖤", "color": "#FE2C55",
        "brand_rgb": RGBColor(0xFE, 0x2C, 0x55),
        "brand_hex": "FE2C55",
        "sub_products": ["放心借", "抖音月付", "DOU分期", "豆包AI", "抖音钱包", "抖音团购", "TikTok Shop"],
        "relevance_keywords": ["抖音", "豆包", "放心借", "团购", "直播", "DOU分期", "抖币", "月付",
                               "字节", "TikTok", "退款", "封号"],
    },
    "美团支付": {
        "icon": "💛", "color": "#FFC300",
        "brand_rgb": RGBColor(0xFF, 0xC3, 0x00),
        "brand_hex": "FFC300",
        "sub_products": ["美团月付", "美团买单", "美团外卖", "美团闪购", "大众点评", "美团借钱", "KeeTa"],
        "relevance_keywords": ["美团", "大众点评", "外卖", "团购", "月付", "闪购", "买单", "配送",
                               "骑手", "无人配送", "KeeTa", "退款", "客服"],
    },
    "京东支付": {
        "icon": "💎", "color": "#E4393C",
        "brand_rgb": RGBColor(0xE4, 0x39, 0x3C),
        "brand_hex": "E4393C",
        "sub_products": ["京东白条", "京东金条", "京东金融", "京东小金库", "京东闪付", "京东钱包"],
        "relevance_keywords": ["京东", "白条", "金条", "小金库", "京东金融", "京东科技", "闪付",
                               "购物卡", "分期", "618", "PLUS", "退款"],
    },
    "云闪付": {
        "icon": "🔴", "color": "#E60012",
        "brand_rgb": RGBColor(0xE6, 0x00, 0x12),
        "brand_hex": "E60012",
        "sub_products": ["银联支付", "Apple Pay", "华为Pay", "数字人民币", "银联国际", "62节"],
        "relevance_keywords": ["云闪付", "银联", "Apple Pay", "华为Pay", "62节", "NFC", "挥卡",
                               "数字人民币", "公交", "地铁", "境外", "退税"],
    },
}

# 客诉分类关键词规则
COMPLAINT_CATEGORIES = {
    "扣费/自动续费": ["自动续费", "扣费", "乱扣", "免密支付", "偷偷扣", "被扣", "不知情扣费"],
    "风控/冻结/封号": ["风控", "冻结", "封号", "封禁", "限制", "异常", "冻结资金", "解封", "永封"],
    "客服/投诉": ["客服", "投诉", "找不到客服", "打不通", "12315", "人工客服", "投诉无门"],
    "退款问题": ["退款", "退款难", "退款慢", "不退款", "拒绝退款", "退款失败"],
    "贷款/借贷": ["花呗", "借呗", "白条", "金条", "月付", "分期", "放心借", "借钱", "逾期", "利息", "利率", "额度"],
    "安全/诈骗": ["被骗", "诈骗", "盗刷", "安全", "骗子", "套路", "维权", "追回"],
    "提现/手续费": ["提现", "手续费", "提现手续费", "提不出来", "到账慢"],
    "功能问题": ["闪退", "bug", "系统繁忙", "无法支付", "支付失败", "绑卡", "系统错误"],
}


def get_likes(post):
    """安全获取点赞数"""
    try:
        return int(str(post.get('liked_count', '0')).replace(',', ''))
    except (ValueError, TypeError):
        return 0


def get_comments(post):
    """安全获取评论数"""
    try:
        return int(str(post.get('comment_count', '0')).replace(',', ''))
    except (ValueError, TypeError):
        return 0


def get_collected(post):
    """安全获取收藏数"""
    try:
        return int(str(post.get('collected_count', '0')).replace(',', ''))
    except (ValueError, TypeError):
        return 0


def get_interactions(post):
    """获取总互动量"""
    return get_likes(post) + get_comments(post) + get_collected(post)


def classify_complaint(title, desc=""):
    """分类客诉帖子"""
    text = f"{title} {desc}".lower()
    matched = []
    for cat, keywords in COMPLAINT_CATEGORIES.items():
        if any(kw in text for kw in keywords):
            matched.append(cat)
    if not matched:
        matched.append("其他问题")
    return matched


def format_number(n):
    """格式化数字"""
    if n >= 10000:
        return f"{n / 10000:.1f}万"
    return f"{n:,}"


def load_platform_data(platform, data_file=None):
    """加载并分析指定平台的数据"""
    if data_file is None:
        data_file = os.path.join(BASE_DIR, "real_data.json")

    with open(data_file, 'r', encoding='utf-8') as f:
        raw = json.load(f)

    all_notes = raw.get('notes', raw) if isinstance(raw, dict) else raw
    
    # 按平台过滤
    platform_notes = [n for n in all_notes if n.get('platform') == platform]
    
    today = date.today()
    report_year = today.year
    report_month = today.month
    report_date_str = today.strftime('%Y年%m月%d日').replace('年0', '年').replace('月0', '月')

    # 客诉帖子
    complaint_notes = []
    for n in platform_notes:
        cats = n.get('categories', [])
        cat = n.get('category', '')
        if '客诉' in cats or cat == '客诉':
            complaint_notes.append(n)

    # 总互动量
    total_interactions = sum(get_interactions(p) for p in platform_notes)
    complaint_interactions = sum(get_interactions(p) for p in complaint_notes)

    # 情绪分析
    sentiment_all = Counter(n.get('sentiment', 'neutral') for n in platform_notes)
    sentiment_complaints = Counter(n.get('sentiment', 'neutral') for n in complaint_notes)

    # 月度统计 - 全部
    monthly_all = Counter()
    for n in platform_notes:
        t = n.get('time', '')[:7]
        if t:
            monthly_all[t] += 1

    # 月度统计 - 客诉
    monthly_complaints = Counter()
    for n in complaint_notes:
        t = n.get('time', '')[:7]
        if t:
            monthly_complaints[t] += 1

    # 分类统计
    category_all = Counter()
    for n in platform_notes:
        cat = n.get('category', '其他')
        category_all[cat] += 1

    # 客诉细分
    complaint_subcats = Counter()
    complaint_subcat_posts = defaultdict(list)
    for n in complaint_notes:
        title = n.get('title', '')
        desc = n.get('desc', '')
        subcats = classify_complaint(title, desc)
        for sc in subcats:
            complaint_subcats[sc] += 1
            complaint_subcat_posts[sc].append(n)

    # 关键词/话题统计
    tag_counts = Counter()
    for n in complaint_notes:
        for tag in n.get('keyword_tags', []):
            tag_counts[tag] += 1

    # Top 帖子（按互动量）
    top_complaints = sorted(complaint_notes, key=get_interactions, reverse=True)[:15]
    top_all = sorted(platform_notes, key=get_interactions, reverse=True)[:10]

    # 时间范围
    all_times = sorted([n.get('time', '')[:10] for n in platform_notes if n.get('time', '')])
    earliest = all_times[0] if all_times else '未知'
    latest = all_times[-1] if all_times else '未知'

    # 当年帖子
    year_prefix = str(report_year)
    posts_year = sorted(
        [n for n in complaint_notes if n.get('time', '').startswith(year_prefix)],
        key=lambda x: x.get('time', '')
    )

    # 当月帖子
    month_prefix = f"{report_year}-{report_month:02d}"
    posts_month = [n for n in complaint_notes if n.get('time', '').startswith(month_prefix)]

    return {
        'platform': platform,
        'all_notes': platform_notes,
        'complaint_notes': complaint_notes,
        'total_count': len(platform_notes),
        'complaint_count': len(complaint_notes),
        'total_interactions': total_interactions,
        'complaint_interactions': complaint_interactions,
        'sentiment_all': sentiment_all,
        'sentiment_complaints': sentiment_complaints,
        'monthly_all': monthly_all,
        'monthly_complaints': monthly_complaints,
        'category_all': category_all,
        'complaint_subcats': complaint_subcats,
        'complaint_subcat_posts': complaint_subcat_posts,
        'tag_counts': tag_counts,
        'top_complaints': top_complaints,
        'top_all': top_all,
        'earliest': earliest,
        'latest': latest,
        'posts_year': posts_year,
        'posts_month': posts_month,
        'report_year': report_year,
        'report_month': report_month,
        'report_date_str': report_date_str,
        'today': today,
    }


def generate_report(platform, year=None, month=None, data_file=None, output_dir=None):
    """为指定平台生成客诉分析报告 Word 文档"""
    if platform not in PLATFORM_CONFIG:
        print(f"❌ 不支持的平台: {platform}")
        print(f"   支持的平台: {', '.join(PLATFORM_CONFIG.keys())}")
        return None

    config = PLATFORM_CONFIG[platform]
    stats = load_platform_data(platform, data_file)
    today = stats['today']
    report_year = year or stats['report_year']
    report_month = month or stats['report_month']
    report_date_str = stats['report_date_str']

    complaint_count = stats['complaint_count']
    total_count = stats['total_count']
    complaint_notes = stats['complaint_notes']
    sentiment_complaints = stats['sentiment_complaints']
    monthly_complaints = stats['monthly_complaints']
    monthly_all = stats['monthly_all']
    complaint_subcats = stats['complaint_subcats']
    complaint_subcat_posts = stats['complaint_subcat_posts']
    tag_counts = stats['tag_counts']
    top_complaints = stats['top_complaints']
    posts_year = stats['posts_year']
    category_all = stats['category_all']
    total_interactions = stats['total_interactions']
    complaint_interactions = stats['complaint_interactions']
    earliest = stats['earliest']
    latest = stats['latest']

    # 格式化
    def format_date_cn(date_str):
        try:
            d = datetime.strptime(date_str, '%Y-%m-%d')
            return f"{d.year}年{d.month}月{d.day}日"
        except Exception:
            return date_str

    interaction_str = format_number(total_interactions) + "次"
    data_range_str = f"{format_date_cn(earliest)} — {format_date_cn(latest)}"
    complaint_pct = f"{complaint_count / total_count * 100:.1f}%" if total_count > 0 else "0%"

    # ==================== 创建文档 ====================
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing_rule = None
    style.paragraph_format.line_spacing = 1.15

    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()

    # 主标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = Pt(42)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run(f"{platform}客诉分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.line_spacing = Pt(22)
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("基于小红书舆情数据的客诉话题分析")
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_horizontal_line(doc)

    # 信息区域
    info_items = [
        ("数据来源", "小红书舆情爬取数据"),
        ("分析时间", report_date_str),
        ("数据范围", data_range_str),
        ("总样本量", f"{total_count}条{platform}相关帖子"),
        ("客诉帖子", f"{complaint_count}条（占比{complaint_pct}）"),
        ("总互动量", interaction_str),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}：")
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_label.font.name = '微软雅黑'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)
        run_val.font.name = '微软雅黑'
        run_val._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_val.font.color.rgb = config['brand_rgb']
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading_styled(doc, "目  录", level=1)
    add_horizontal_line(doc)

    toc_items = [
        "一、总结",
        "二、客诉分类统计与详细分析",
        "三、情绪分析",
        "四、月度趋势变化",
        "五、热点话题",
        "六、建议与改进方向",
        "七、结论",
        f"附录：{report_year}年客诉帖子清单",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item.strip())
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ==================== 一、总结 ====================
    add_heading_styled(doc, "一、总结", level=1)
    add_horizontal_line(doc)

    # 总体概览
    neg_count = sentiment_complaints.get('negative', 0)
    neg_pct = f"{neg_count / complaint_count * 100:.0f}%" if complaint_count > 0 else "0%"

    add_callout_box(doc,
        f"截至{report_date_str}，{platform}在小红书上的相关帖子共计{total_count}条，"
        f"其中客诉类帖子{complaint_count}条（占比{complaint_pct}），"
        f"客诉帖子的总互动量超过{format_number(complaint_interactions)}次。",
        bg_color="E8F0FE", border_color=config['brand_hex']
    )

    # 话题分类概览表
    add_paragraph_styled(doc, "话题分类概览：", bold=True, font_size=11, space_after=8)
    cat_rows = []
    for cat, cnt in category_all.most_common():
        pct = f"{cnt / total_count * 100:.1f}%" if total_count > 0 else "0%"
        cat_rows.append([cat, str(cnt), pct])
    create_styled_table(doc,
        ["话题类型", "帖子数", "占比"],
        cat_rows,
        col_widths=[5, 3, 3]
    )

    doc.add_paragraph()

    # 客诉细分概览表
    add_paragraph_styled(doc, "客诉问题细分概览：", bold=True, font_size=11, space_after=8)
    subcat_rows = []
    rank = 0
    for subcat, cnt in complaint_subcats.most_common():
        rank += 1
        pct = f"{cnt / complaint_count * 100:.1f}%" if complaint_count > 0 else "0%"
        # 找最高赞帖子
        posts_in_cat = complaint_subcat_posts[subcat]
        top_post = max(posts_in_cat, key=get_likes) if posts_in_cat else None
        top_likes = format_number(get_likes(top_post)) if top_post else "-"
        subcat_rows.append([str(rank), subcat, str(cnt), pct, top_likes])
    create_styled_table(doc,
        ["排名", "问题类型", "帖子数", "占比", "最高单帖赞"],
        subcat_rows,
        col_widths=[1.5, 4, 2, 2, 2.5]
    )

    doc.add_paragraph()

    # 关键发现
    add_paragraph_styled(doc, "关键发现：", bold=True, font_size=11)

    # 动态生成关键发现
    if complaint_subcats:
        top_subcat = complaint_subcats.most_common(1)[0]
        top_subcat_name = top_subcat[0]
        top_subcat_cnt = top_subcat[1]
        top_subcat_pct = f"{top_subcat_cnt / complaint_count * 100:.0f}%" if complaint_count > 0 else "N/A"
        add_bullet_point(doc,
            f"是最主要的客诉类型，共{top_subcat_cnt}条帖子，占全部客诉的{top_subcat_pct}。",
            bold_prefix=f'\u201c{top_subcat_name}\u201d')

    if top_complaints:
        top1 = top_complaints[0]
        top1_title = top1.get('title', '')[:30]
        add_bullet_point(doc,
            f"（{format_number(get_interactions(top1))}次互动），反映了用户最关注的痛点。",
            bold_prefix=f'最高互动帖子：\u201c{top1_title}\u2026\u201d')

    neg_ratio = neg_count / complaint_count * 100 if complaint_count > 0 else 0
    if neg_ratio > 50:
        add_bullet_point(doc,
            f"，{neg_pct}的客诉帖子带有负面情绪，说明用户不满情绪强烈。",
            bold_prefix="负面情绪占比高")
    elif neg_ratio > 30:
        add_bullet_point(doc,
            f"，{neg_pct}的客诉帖子带有负面情绪。",
            bold_prefix="负面情绪占比较高")

    # 月度趋势简述
    sorted_months = sorted(monthly_complaints.items())
    if len(sorted_months) >= 2:
        last_month_key = sorted_months[-1]
        prev_month_key = sorted_months[-2]
        if last_month_key[1] > prev_month_key[1]:
            add_bullet_point(doc,
                f"最近月度客诉量呈上升趋势，{last_month_key[0]}有{last_month_key[1]}条，"
                f"较上月（{prev_month_key[1]}条）增长{last_month_key[1] - prev_month_key[1]}条。",
                bold_prefix="趋势上升：")

    doc.add_page_break()

    # ==================== 二、客诉分类统计与详细分析 ====================
    add_heading_styled(doc, "二、客诉分类统计与详细分析", level=1)
    add_horizontal_line(doc)

    for idx, (subcat, cnt) in enumerate(complaint_subcats.most_common()):
        pct = f"{cnt / complaint_count * 100:.1f}%" if complaint_count > 0 else "0%"
        add_heading_styled(doc, f"2.{idx + 1} {subcat}（{cnt}条，占比{pct}）", level=2)

        # 对该类客诉的描述
        posts_in_cat = complaint_subcat_posts[subcat]

        # 情绪分析
        cat_sentiments = Counter(p.get('sentiment', 'neutral') for p in posts_in_cat)
        cat_neg = cat_sentiments.get('negative', 0)
        cat_neg_pct = f"{cat_neg / cnt * 100:.0f}%" if cnt > 0 else "0%"
        
        add_callout_box(doc,
            f"此类客诉共{cnt}条帖子，负面情绪占比{cat_neg_pct}。"
            f"总互动量{format_number(sum(get_interactions(p) for p in posts_in_cat))}次。")

        # 典型帖子表格（取top 8）
        top_cat_posts = sorted(posts_in_cat, key=get_interactions, reverse=True)[:8]
        if top_cat_posts:
            add_paragraph_styled(doc, "典型帖子：", bold=True, font_size=10.5, space_after=4)
            table_rows = []
            for p in top_cat_posts:
                t = p.get('time', '')[:10]
                title = p.get('title', '')[:40]
                likes = format_number(get_likes(p))
                comments = format_number(get_comments(p))
                table_rows.append([t, title, likes, comments])
            create_styled_table(doc,
                ["时间", "标题", "点赞", "评论"],
                table_rows,
                col_widths=[2.5, 8, 2, 2]
            )

        doc.add_paragraph()

        # 用户核心诉求（基于标题关键词提取）
        add_paragraph_styled(doc, "用户核心诉求：", bold=True, font_size=10.5)
        # 从标题中提取高频词
        title_words = Counter()
        for p in posts_in_cat:
            title = p.get('title', '')
            for kw in COMPLAINT_CATEGORIES.get(subcat, []):
                if kw in title:
                    title_words[kw] += 1
        
        # 列出代表性标题
        shown = 0
        for p in sorted(posts_in_cat, key=get_likes, reverse=True)[:5]:
            title = p.get('title', '')
            if title and shown < 5:
                add_bullet_point(doc, f'（{format_number(get_likes(p))}赞）', bold_prefix=f'"{title[:50]}"')
                shown += 1

        if idx < len(complaint_subcats) - 1:
            add_horizontal_line(doc)

    doc.add_page_break()

    # ==================== 三、情绪分析 ====================
    add_heading_styled(doc, "三、情绪分析", level=1)
    add_horizontal_line(doc)

    add_paragraph_styled(doc, "全部帖子情绪分布：", bold=True, font_size=11)
    sentiment_labels = {'positive': '正面', 'negative': '负面', 'neutral': '中性'}
    sentiment_all = stats['sentiment_all']
    sent_rows = []
    for key in ['positive', 'neutral', 'negative']:
        cnt = sentiment_all.get(key, 0)
        pct = f"{cnt / total_count * 100:.1f}%" if total_count > 0 else "0%"
        bar = "█" * max(1, int(cnt / total_count * 30)) if total_count > 0 else "█"
        sent_rows.append([sentiment_labels.get(key, key), str(cnt), pct, bar])
    create_styled_table(doc,
        ["情绪", "数量", "占比", "分布"],
        sent_rows,
        col_widths=[2, 2, 2, 10]
    )

    doc.add_paragraph()

    add_paragraph_styled(doc, "客诉帖子情绪分布：", bold=True, font_size=11)
    sent_c_rows = []
    for key in ['positive', 'neutral', 'negative']:
        cnt = sentiment_complaints.get(key, 0)
        pct = f"{cnt / complaint_count * 100:.1f}%" if complaint_count > 0 else "0%"
        bar = "█" * max(1, int(cnt / complaint_count * 30)) if complaint_count > 0 else "█"
        sent_c_rows.append([sentiment_labels.get(key, key), str(cnt), pct, bar])
    create_styled_table(doc,
        ["情绪", "数量", "占比", "分布"],
        sent_c_rows,
        col_widths=[2, 2, 2, 10]
    )

    doc.add_paragraph()

    # 情绪分析总结
    total_neg = sentiment_all.get('negative', 0)
    total_neg_pct = total_neg / total_count * 100 if total_count > 0 else 0
    comp_neg_pct_val = neg_count / complaint_count * 100 if complaint_count > 0 else 0

    add_callout_box(doc,
        f"整体来看，{platform}在小红书上的负面情绪帖子占比{total_neg_pct:.1f}%，"
        f"而在客诉帖子中负面情绪占比上升至{comp_neg_pct_val:.1f}%，"
        f"说明客诉场景下用户的不满情绪更为集中。",
        bg_color="FFF3E0", border_color="E65100"
    )

    doc.add_page_break()

    # ==================== 四、月度趋势变化 ====================
    add_heading_styled(doc, "四、月度趋势变化", level=1)
    add_horizontal_line(doc)

    add_paragraph_styled(doc, "全部帖子月度趋势：", bold=True, font_size=11)

    # 生成月度趋势表
    all_months_sorted = sorted(monthly_all.keys())
    if all_months_sorted:
        max_count = max(monthly_all.values()) if monthly_all else 1
        trend_rows = []
        for mk in all_months_sorted:
            cnt = monthly_all[mk]
            parts = mk.split('-')
            label = f"{parts[0]}年{int(parts[1])}月"
            bar_len = max(1, int(cnt / max_count * 30))
            trend_rows.append([label, str(cnt), "█" * bar_len])
        create_styled_table(doc,
            ["月份", "数量", "趋势"],
            trend_rows,
            col_widths=[3, 2, 11]
        )

    doc.add_paragraph()

    add_paragraph_styled(doc, "客诉帖子月度趋势：", bold=True, font_size=11)

    complaint_months_sorted = sorted(monthly_complaints.keys())
    if complaint_months_sorted:
        max_c_count = max(monthly_complaints.values()) if monthly_complaints else 1
        trend_c_rows = []
        for mk in complaint_months_sorted:
            cnt = monthly_complaints[mk]
            parts = mk.split('-')
            label = f"{parts[0]}年{int(parts[1])}月"
            bar_len = max(1, int(cnt / max_c_count * 30))
            marker = " ⬆" if mk == complaint_months_sorted[-1] and cnt == max_c_count else ""
            trend_c_rows.append([label, f"{cnt}{marker}", "█" * bar_len])
        create_styled_table(doc,
            ["月份", "数量", "趋势"],
            trend_c_rows,
            col_widths=[3, 2, 11]
        )

    doc.add_paragraph()

    # 趋势分析总结
    add_paragraph_styled(doc, "趋势分析：", bold=True, font_size=11)
    if len(complaint_months_sorted) >= 3:
        recent3 = complaint_months_sorted[-3:]
        recent3_counts = [monthly_complaints[m] for m in recent3]
        if recent3_counts[-1] > recent3_counts[0]:
            add_bullet_point(doc,
                f"近3个月客诉量呈上升趋势（{' → '.join(str(c) for c in recent3_counts)}条），需要引起关注。",
                bold_prefix="上升趋势：")
        elif recent3_counts[-1] < recent3_counts[0]:
            add_bullet_point(doc,
                f"近3个月客诉量呈下降趋势（{' → '.join(str(c) for c in recent3_counts)}条），改善初见成效。",
                bold_prefix="下降趋势：")
        else:
            add_bullet_point(doc,
                f"近3个月客诉量相对稳定（{' → '.join(str(c) for c in recent3_counts)}条）。",
                bold_prefix="趋于稳定：")

    doc.add_page_break()

    # ==================== 五、热点话题 ====================
    add_heading_styled(doc, "五、热点话题", level=1)
    add_horizontal_line(doc)

    add_paragraph_styled(doc, "客诉帖子高频标签 Top 15：", bold=True, font_size=11)
    tag_rows = []
    for rank_i, (tag, cnt) in enumerate(tag_counts.most_common(15)):
        tag_rows.append([str(rank_i + 1), tag, str(cnt)])
    if tag_rows:
        create_styled_table(doc,
            ["排名", "话题标签", "出现次数"],
            tag_rows,
            col_widths=[2, 8, 3]
        )

    doc.add_paragraph()

    add_paragraph_styled(doc, "最高互动量客诉帖子 Top 10：", bold=True, font_size=11)
    hot_rows = []
    for p in top_complaints[:10]:
        t = p.get('time', '')[:10]
        title = p.get('title', '')[:35]
        inter = format_number(get_interactions(p))
        likes = format_number(get_likes(p))
        sent = sentiment_labels.get(p.get('sentiment', 'neutral'), '中性')
        hot_rows.append([t, title, likes, inter, sent])
    if hot_rows:
        create_styled_table(doc,
            ["时间", "标题", "点赞", "互动量", "情绪"],
            hot_rows,
            col_widths=[2.5, 6, 2, 2, 1.5]
        )

    doc.add_page_break()

    # ==================== 六、建议与改进方向 ====================
    add_heading_styled(doc, "六、建议与改进方向", level=1)
    add_horizontal_line(doc)

    add_paragraph_styled(doc,
        f"基于对{platform}客诉数据的分析，以下建议按问题严重性和影响面排列：",
        font_size=10.5, space_after=8)

    # 为每个客诉子类生成建议
    priority_colors = ["🔴", "🔴", "🟠", "🟠", "🟡", "🟡", "🟡", "🟢"]
    priority_labels = ["最高", "最高", "高", "高", "中", "中", "中", "较低"]

    for idx, (subcat, cnt) in enumerate(complaint_subcats.most_common()):
        priority_idx = min(idx, len(priority_colors) - 1)
        priority = f"{priority_colors[priority_idx]} {priority_labels[priority_idx]}"
        add_heading_styled(doc,
            f"6.{idx + 1} {subcat}（优先级：{priority}）", level=2)

        # 根据客诉类型生成通用建议
        suggestions = _get_suggestions_for_category(subcat, platform)
        if suggestions:
            create_styled_table(doc,
                ["建议", "具体方案"],
                suggestions,
                col_widths=[4, 12]
            )
        doc.add_paragraph()

    doc.add_page_break()

    # ==================== 七、结论 ====================
    add_heading_styled(doc, "七、结论", level=1)
    add_horizontal_line(doc)

    # 动态生成结论
    top3_subcats = complaint_subcats.most_common(3)
    if top3_subcats:
        top3_names = "\u3001".join(f'\u201c{sc[0]}\u201d' for sc in top3_subcats)
        top3_total = sum(sc[1] for sc in top3_subcats)
        top3_pct = f"{top3_total / complaint_count * 100:.0f}%" if complaint_count > 0 else "N/A"
        add_paragraph_styled(doc,
            f"{platform}面临的客诉问题主要集中在{top3_names}三个领域。"
            f"这三类问题合计占全部客诉的{top3_pct}，"
            f"是最需要优先解决的用户痛点。",
            bold=True, font_size=11)

    doc.add_paragraph()

    add_callout_box(doc,
        f"核心改进方向：{platform}需要在产品透明度、客服响应速度和用户权益保障三个维度持续优化。"
        f"让用户看得见（信息透明）、管得了（自主操控）、找得到（客服可达）"
        f"——是提升用户满意度的关键。",
        bg_color="E8F5E9", border_color="1B5E20"
    )

    doc.add_paragraph()

    # ==================== 附录：帖子清单 ====================
    doc.add_page_break()
    add_heading_styled(doc, f"附录：{report_year}年客诉帖子清单", level=1)
    add_horizontal_line(doc)

    add_paragraph_styled(doc,
        f"以下为{report_year}年期间，小红书上{platform}相关客诉帖子的完整清单（共{len(posts_year)}条），"
        "按时间顺序排列。每条帖子均附有原始链接，可点击查看原文。",
        font_size=10, space_after=10)

    # 按月份分组
    months_map = defaultdict(list)
    for post in posts_year:
        month_key = post.get('time', '')[:7]
        months_map[month_key].append(post)

    for month_key in sorted(months_map.keys()):
        month_posts = months_map[month_key]
        parts = month_key.split('-')
        label = f"{parts[0]}年{int(parts[1])}月"

        add_heading_styled(doc, f"{label}（{len(month_posts)}条）", level=2)

        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, h in enumerate(["序号", "标题", "日期"]):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, config['brand_hex'])

        # 数据行
        for idx_r, post in enumerate(month_posts):
            row = table.add_row()

            # 序号
            cell0 = row.cells[0]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(str(idx_r + 1))
            run0.font.size = Pt(8)
            run0.font.name = '微软雅黑'
            run0._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 标题（带超链接）
            cell1 = row.cells[1]
            p1 = cell1.paragraphs[0]
            title_text = post.get('title', '')
            link = post.get('link', '')
            if link:
                hyperlink = OxmlElement('w:hyperlink')
                r_id = doc.part.relate_to(link,
                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                    is_external=True)
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), '微软雅黑')
                rFonts.set(qn('w:eastAsia'), '微软雅黑')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '16')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '16')
                rPr.append(szCs)
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), config['brand_hex'])
                rPr.append(color_el)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                new_run.append(rPr)
                new_run.text = title_text
                hyperlink.append(new_run)
                p1._element.append(hyperlink)
            else:
                run1 = p1.add_run(title_text)
                run1.font.size = Pt(8)
                run1.font.name = '微软雅黑'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 日期
            cell2 = row.cells[2]
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(post.get('time', '')[:10])
            run2.font.size = Pt(8)
            run2.font.name = '微软雅黑'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            if idx_r % 2 == 1:
                for c in row.cells:
                    set_cell_shading(c, "E8EDF3")

        # 列宽
        for row in table.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(11)
            row.cells[2].width = Cm(3.5)

        # 表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        doc.add_paragraph()

    add_paragraph_styled(doc,
        f"共计 {len(posts_year)} 条帖子。所有链接均指向小红书原帖，可直接点击查看。",
        italic=True, color=RGBColor(0x66, 0x66, 0x66), font_size=9, space_after=12)

    # 页脚
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_line.add_run("— 报告结束 —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_time_str = today.strftime('%Y-%m-%d')
    run = footer_info.add_run(f"报告生成时间：{gen_time_str} | 数据源：小红书舆情爬取系统")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ==================== 保存 ====================
    if output_dir is None:
        output_dir = BASE_DIR
    os.makedirs(output_dir, exist_ok=True)

    output_filename = f"{platform}客诉分析报告_{report_year}年{report_month:02d}月.docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    print(f"✅ {platform}报告已生成：{output_path}")
    return output_path


def _get_suggestions_for_category(subcat, platform):
    """根据客诉类型生成通用改进建议"""
    suggestions_map = {
        "扣费/自动续费": [
            ["优化入口可发现性", f"在{platform}主页面增加醒目的「扣费管理」入口"],
            ["主动提醒机制", "每月推送一次「本月自动扣费汇总」通知"],
            ["一键管理页面", "建立统一的自动续费/免密支付管理页面"],
            ["扣费前确认", "对长期未使用但仍自动续费的服务，扣费前发送确认"],
        ],
        "风控/冻结/封号": [
            ["降低误伤率", "优化风控模型，对正常用户放宽阈值"],
            ["缩短审核周期", "为低风险case提供快速解封通道"],
            ["透明化风控原因", "告诉用户具体被风控的原因"],
            ["自助解封通道", "提供人脸+短信验证的快速自助解封"],
        ],
        "客服/投诉": [
            ["人工客服入口外露", "在问题场景中直接展示人工客服按钮"],
            ["AI客服转人工", "当AI无法解决时快速转接人工"],
            ["投诉工单跟踪", "建立可追踪的工单系统"],
            ["响应时效承诺", "明确投诉处理时限并主动更新进度"],
        ],
        "退款问题": [
            ["加速退款处理", "承诺3个工作日内完成退款并到账"],
            ["退款进度透明", "提供实时退款进度查询"],
            ["简化退款流程", "减少退款所需的操作步骤"],
            ["退款原因说明", "退款拒绝时给出具体原因和解决方案"],
        ],
        "贷款/借贷": [
            ["利率透明化", "在产品页面显著展示年化利率"],
            ["还款提醒", "提前多天提醒还款日和金额"],
            ["额度管理", "额度变化时及时通知并说明原因"],
            ["逾期处理", "提供灵活的逾期还款方案"],
        ],
        "安全/诈骗": [
            ["转账安全提醒", "对大额转账增加二次确认"],
            ["资金追回机制", "被骗后提供便捷的冻结和追回通道"],
            ["安全教育", "在App内提供防诈骗指引"],
            ["实时风险拦截", "加强对可疑转账的实时检测"],
        ],
        "提现/手续费": [
            ["降低费率", "继续降低提现手续费或增加免费额度"],
            ["费用透明", "在提现页面明确展示费用和免费方案"],
            ["免费替代方案", "提供更多免手续费的提现途径"],
        ],
        "功能问题": [
            ["稳定性优化", "加强技术稳定性，减少闪退和系统错误"],
            ["用户体验优化", "简化操作流程，降低使用门槛"],
            ["错误提示改善", "提供明确的错误信息和解决方案"],
            ["兼容性测试", "加强多机型、多版本的兼容性测试"],
        ],
        "其他问题": [
            ["用户反馈渠道", "建立便捷的用户反馈收集机制"],
            ["问题分类跟踪", "对新出现的问题及时分类并跟进"],
            ["定期用户调研", "开展定期用户满意度调研"],
        ],
    }
    return suggestions_map.get(subcat, [
        ["收集反馈", "建立该类问题的专项反馈渠道"],
        ["跟踪改进", "设置专项改进指标并定期复盘"],
    ])


def generate_all_platform_reports(year=None, month=None, output_dir=None):
    """批量生成全部5个非微信平台的月报"""
    results = []
    for platform in PLATFORM_CONFIG.keys():
        try:
            path = generate_report(platform, year=year, month=month, output_dir=output_dir)
            if path:
                results.append(path)
        except Exception as e:
            print(f"❌ {platform} 报告生成失败: {e}")
            import traceback
            traceback.print_exc()
    return results


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="通用多平台客诉分析报告生成器")
    parser.add_argument('--platform', type=str, help='平台名称（支付宝/抖音支付/美团支付/京东支付/云闪付）')
    parser.add_argument('--all', action='store_true', help='生成全部5个平台的报告')
    parser.add_argument('--year', type=int, default=None, help='报告年份')
    parser.add_argument('--month', type=int, default=None, help='报告月份')
    parser.add_argument('--output-dir', type=str, default=None, help='输出目录')

    args = parser.parse_args()

    if args.all:
        results = generate_all_platform_reports(year=args.year, month=args.month, output_dir=args.output_dir)
        print(f"\n{'=' * 60}")
        print(f"🎉 批量生成完成！共生成 {len(results)} 份报告")
        for r in results:
            print(f"   📄 {os.path.basename(r)}")
        print(f"{'=' * 60}")
    elif args.platform:
        generate_report(args.platform, year=args.year, month=args.month, output_dir=args.output_dir)
    else:
        print("请指定 --platform 或 --all")
        parser.print_help()
